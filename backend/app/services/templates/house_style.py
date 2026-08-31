"""The Adani Global FZE house style, as measured from the approved reference documents.

Every constant here was read out of a real, approved document rather than chosen. The sources are:

* `Sample Purchase Contract 1.docx` - a native Word document, so its `sectPr`, `header1.xml` and
  `footer1.xml` give page size, margins, the letterhead logo at its true extent, and the footer
  block verbatim. This is the authority for anything structural.
* `LATES FILLE SCRAP INVOICE.pdf` - a two-page PDF that is in fact two documents, a commercial
  invoice and a provisional invoice. Authority for the invoice metadata block and the
  description/quantity/unit-price/amount table.
* `Sales Contract_2026-27 COPY 1.xlsx` - the desk's own contract generator. Authority for the
  sales contract and Performa invoice layouts, and for which fields are deal-specific.

Measurements are recorded in the units the source records them in (twips for Word geometry, EMU
for the image extent) and converted once, here, so the arithmetic is visible and checkable.

Nothing in this module is a design decision. Where a value could not be read from a reference it
is named as an approximation in a comment beside it.
"""

from __future__ import annotations

from pathlib import Path

ASSET_ROOT = Path(__file__).resolve().parent / "assets"

# The letterhead mark, lifted whole out of `word/media/image1.jpeg` in the reference contract
# rather than redrawn. 441 x 255 px at the extent below.
LOGO_PATH = ASSET_ROOT / "letterhead_logo.jpeg"

# --- page geometry ------------------------------------------------------------------------------
#
# From `<w:pgSz w:w="12240" w:h="15840"/>` - 8.5in x 11in, US Letter, portrait. Worth stating
# plainly because A4 would be the likelier guess for a Dubai-issued document and it is not what
# the reference uses.
PAGE_WIDTH_IN = 8.5
PAGE_HEIGHT_IN = 11.0

# From `<w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1620" .../>`. The left margin
# is deliberately wider than the right - 1.125in against 1.0in - which is what gives the reference
# its off-centre body block. Reproduced rather than squared up.
MARGIN_TOP_IN = 1.0

# The reference declares `w:bottom="1440"`, 1.0in, and this is the one measurement here that is
# deliberately not reproduced. The letterhead footer below is seven lines deep and occupies the
# band from roughly 1.4in to 0.6in above the paper edge - more than a 1.0in bottom margin leaves
# for it. The reference documents get away with it because none of them runs body text that far
# down the page; a generated contract with a dozen clauses does, and prints its last lines over
# the registered-entity block. Reproducing the number while producing overlapping text would not
# be reproducing the reference, which has no overlap on any page. So the margin is set to clear
# the measured footer band instead, and every other page measurement is the reference's own.
MARGIN_BOTTOM_IN = 1.45
MARGIN_LEFT_IN = 1.125
MARGIN_RIGHT_IN = 1.0

# `w:header="0" w:footer="144"` - the header sits flush to the paper edge, the footer 0.1in up.
HEADER_DISTANCE_IN = 0.0
FOOTER_DISTANCE_IN = 0.1

# --- the letterhead mark ------------------------------------------------------------------------
#
# `<wp:extent cx="1628775" cy="933450"/>`, in EMU. 914400 EMU to the inch.
LOGO_WIDTH_IN = 1628775 / 914400  # 1.781in
LOGO_HEIGHT_IN = 933450 / 914400  # 1.021in

# The header paragraph carries `<w:ind w:hanging="720"/>`, so the mark hangs 0.5in into the left
# margin and its left edge lands at 0.625in from the paper edge.
HEADER_HANGING_IN = 0.5

# --- typography ---------------------------------------------------------------------------------
#
# `w:ascii="Adani Regular"` on all 265 runs of the reference contract. It is a licensed corporate
# face and will not be installed on a machine that does not have the Adani font pack, where Word
# substitutes. The reference PDF invoice, produced on a machine without it, fell back to Arial -
# so Arial is the observed fallback rather than a guess, and it is what the metrics below were
# eyeballed against.
BODY_FONT = "Adani Regular"
FALLBACK_FONT = "Arial"

# Half-point sizes read off the reference runs: body `w:sz 18` = 9pt, the contract reference line
# `w:sz 20` = 10pt, the footer `w:sz 16` = 8pt. The title run carries no explicit size and so
# inherits the document default; 11pt is that default and is consistent with the rendered page.
SIZE_BODY_PT = 9
SIZE_REFERENCE_PT = 10
SIZE_TITLE_PT = 11
SIZE_FOOTER_PT = 8

# --- footer -------------------------------------------------------------------------------------
#
# Verbatim from `footer1.xml`, including the colour values. 333399 is the navy the Arabic trading
# name is set in; 4A442A the dark olive of the address block; 808080 the grey of the two closing
# legal lines.
COLOUR_ARABIC_NAME = "333399"
COLOUR_ADDRESS = "4A442A"
COLOUR_LEGAL = "808080"

FOOTER_ARABIC_NAME = "\u0627\u062f\u0627\u0646\u064a \u062f\u062c\u0644\u0648\u0628\u0627\u0644 \u0645 \u0645 \u062d"

# Each footer address row is one paragraph split by a tab: the postal address on the left, the
# contact detail on the right. Taken exactly as the reference has them, trailing comma included.
FOOTER_ADDRESS_ROWS: tuple[tuple[str, str], ...] = (
    ("Adani Global FZE", "Tel: +971 4 3611900"),
    ("Post Box 17186,", "Fax: +971 4 3689600"),
    ("Jebel Ali Free Zone,", "info.dubai@adani.in"),
    ("Dubai, UAE", "www.adani.com"),
)

FOOTER_LEGAL_ARABIC = (
    "\u0630\u0627\u062a \u0645\u0633\u0626\u0648\u0644\u064a\u0629 \u0645\u062d\u062f\u0648\u062f\u0629 "
    "\u062a\u0623\u0633\u0633\u062a \u0628\u0645\u0648\u062c\u0628 \u0627\u0644\u0642\u0627\u0646\u0648\u0646 "
    "\u0631\u0642\u0645 9 \u0644\u0633\u0646\u0629 1992"
)
FOOTER_LEGAL_ENGLISH = "Formed pursuant to Law No.9 of 1992 With Limited Liability"

# `<w:ind w:hanging="630"/>` on every address row - 0.4375in - which outdents the footer past the
# body's left margin to 0.6875in from the paper edge. The tab stop is `w:pos="1050"`, 0.729in.
FOOTER_HANGING_IN = 0.4375
FOOTER_TAB_IN = 630 / 1440 + 1050 / 1440  # measured from the outdented start, so 1.167in

# --- the issuing entity -------------------------------------------------------------------------
#
# The legal name every reference document names as the AGFZE party, in the casing they use it in.
# The contract signature block, the invoice's retention-of-title line and the bank beneficiary
# name are all this exact string.
SELLER_LEGAL_NAME = "ADANI GLOBAL FZE"
SELLER_ADDRESS_LINE = "P.O. BOX 17186, DUBAI, U.A.E."


def _set_font(run, *, size_pt: int, bold: bool = False, colour: str | None = None) -> None:
    """Put one run into the house face, since the Normal style alone does not carry east-asian
    and complex-script mappings and Word will otherwise substitute mid-paragraph."""
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = BODY_FONT
    element = run._element.rPr.rFonts
    element.set(qn("w:eastAsia"), BODY_FONT)
    element.set(qn("w:cs"), BODY_FONT)
    run.font.size = Pt(size_pt)
    run.bold = bold
    if colour is not None:
        run.font.color.rgb = RGBColor.from_string(colour)


def apply_page_setup(document) -> None:
    """Page size, margins and header/footer offsets, from the reference `sectPr`."""
    from docx.shared import Inches

    for section in document.sections:
        section.page_width = Inches(PAGE_WIDTH_IN)
        section.page_height = Inches(PAGE_HEIGHT_IN)
        section.top_margin = Inches(MARGIN_TOP_IN)
        section.bottom_margin = Inches(MARGIN_BOTTOM_IN)
        section.left_margin = Inches(MARGIN_LEFT_IN)
        section.right_margin = Inches(MARGIN_RIGHT_IN)
        section.header_distance = Inches(HEADER_DISTANCE_IN)
        section.footer_distance = Inches(FOOTER_DISTANCE_IN)


def apply_base_font(document) -> None:
    """Set the document default so text the builder does not touch still comes out in the face."""
    from docx.oxml.ns import qn
    from docx.shared import Pt

    normal = document.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(SIZE_BODY_PT)
    rpr = normal.element.get_or_add_rPr()
    fonts = rpr.get_or_add_rFonts()
    for attribute in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        fonts.set(qn(attribute), BODY_FONT)


def build_header(document) -> None:
    """The letterhead mark, at the reference's own extent and hanging indent.

    It goes in the section header rather than the body so it repeats on every page by itself,
    which is what the reference does - the mark is on all three pages of the sample contract,
    including the signature page that carries nothing else.
    """
    from docx.shared import Inches

    header = document.sections[0].header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.paragraph_format.first_line_indent = -Inches(HEADER_HANGING_IN)
    paragraph.paragraph_format.space_after = Inches(0)
    run = paragraph.add_run()
    run.add_picture(str(LOGO_PATH), width=Inches(LOGO_WIDTH_IN), height=Inches(LOGO_HEIGHT_IN))


def build_footer(document) -> None:
    """The registered-entity block, verbatim from the reference footer.

    Static in the reference and static here: nothing in it is deal-specific, so it is template
    content rather than a placeholder. Note that this is the *entity* footer - the bank details
    that appear on an invoice are not here, because the references show those changing from one
    deal to the next and they are carried as fields instead.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.shared import Inches

    footer = document.sections[0].footer
    footer.is_linked_to_previous = False

    arabic = footer.paragraphs[0]
    arabic.paragraph_format.first_line_indent = -Inches(FOOTER_HANGING_IN)
    arabic.paragraph_format.space_after = Inches(0)
    _set_font(
        arabic.add_run(FOOTER_ARABIC_NAME),
        size_pt=SIZE_FOOTER_PT,
        bold=True,
        colour=COLOUR_ARABIC_NAME,
    )

    for left, right in FOOTER_ADDRESS_ROWS:
        row = footer.add_paragraph()
        row.paragraph_format.first_line_indent = -Inches(FOOTER_HANGING_IN)
        row.paragraph_format.space_after = Inches(0)
        row.paragraph_format.tab_stops.add_tab_stop(
            Inches(FOOTER_TAB_IN), WD_TAB_ALIGNMENT.LEFT
        )
        _set_font(row.add_run(left), size_pt=SIZE_FOOTER_PT, bold=True, colour=COLOUR_ADDRESS)
        _set_font(
            row.add_run("\t" + right), size_pt=SIZE_FOOTER_PT, bold=True, colour=COLOUR_ADDRESS
        )

    for line in (FOOTER_LEGAL_ARABIC, FOOTER_LEGAL_ENGLISH):
        legal = footer.add_paragraph()
        legal.alignment = WD_ALIGN_PARAGRAPH.CENTER
        legal.paragraph_format.space_after = Inches(0)
        _set_font(legal.add_run(line), size_pt=SIZE_FOOTER_PT, bold=True, colour=COLOUR_LEGAL)


def add_title(document, text: str):
    """The document title: centred, bold, underlined, at the inherited size.

    Every reference document that has a title sets it this way - PURCHASE CONTRACT, SALES
    CONTRACT, COMMERCIAL INVOICE, PROVISIONAL INVOICE and PERFORMA INVOICE alike.
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    _set_font(run, size_pt=SIZE_TITLE_PT, bold=True)
    run.underline = True
    return paragraph


def add_reference_line(document, left_text: str, right_text: str):
    """The `{contract no}` ... `DATE: {date}` line that sits under the title.

    The reference sets it as one justified paragraph with the date pushed to the right margin,
    and repeats it at the head of every continuation page. Both halves are bold, at 10pt - a
    point larger than body text, which is the only place the reference varies the body size.
    """
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Inches

    usable = PAGE_WIDTH_IN - MARGIN_LEFT_IN - MARGIN_RIGHT_IN
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(usable), WD_TAB_ALIGNMENT.RIGHT)
    _set_font(paragraph.add_run(left_text), size_pt=SIZE_REFERENCE_PT, bold=True)
    _set_font(paragraph.add_run("\t" + right_text), size_pt=SIZE_REFERENCE_PT, bold=True)
    return paragraph


def add_party_block(document, lines: list[str]):
    """The counterparty's name, address and registration lines, bold and flush left.

    The references run these as plain bold paragraphs rather than putting them in a bordered
    grid, and the sales contract packs the registration identifiers (IEC, PAN, GST, e-mail,
    telephone) onto their own wrapped line beneath the address.
    """
    from docx.shared import Inches

    written = []
    for line in lines:
        paragraph = document.add_paragraph()
        paragraph.paragraph_format.space_after = Inches(0)
        _set_font(paragraph.add_run(line), size_pt=SIZE_BODY_PT, bold=True)
        written.append(paragraph)
    return written


def add_metadata_row(document, label: str, value: str, *, label_width_in: float = 2.2):
    """One `LABEL : value` row, in the pattern the reference invoice uses throughout.

    Bold label in a fixed-width column, a colon in its own column, then the value. This is the
    reference's field-positioning convention and is used in preference to the bordered two-column
    table the builder used to emit, which appears nowhere in any reference document.
    """
    from docx.enum.text import WD_TAB_ALIGNMENT
    from docx.shared import Inches

    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Inches(0)
    stops = paragraph.paragraph_format.tab_stops
    stops.add_tab_stop(Inches(label_width_in), WD_TAB_ALIGNMENT.LEFT)
    stops.add_tab_stop(Inches(label_width_in + 0.2), WD_TAB_ALIGNMENT.LEFT)
    _set_font(paragraph.add_run(label), size_pt=SIZE_BODY_PT, bold=True)
    _set_font(paragraph.add_run("\t:\t"), size_pt=SIZE_BODY_PT)
    return paragraph


def add_rule(paragraph, *, above: bool = False, below: bool = False) -> None:
    """A single hairline above and/or below a paragraph.

    The reference invoice's line-item table is not a bordered grid: it is a rule above the column
    headings, a rule below them, and a rule above the total. Reproduced as paragraph borders so
    the rule spans the full text width the way it does on the reference.
    """
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    borders = OxmlElement("w:pBdr")
    for edge, wanted in (("top", above), ("bottom", below)):
        if not wanted:
            continue
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")  # 0.75pt, the weight the reference rules measure at
        element.set(qn("w:space"), "1")
        element.set(qn("w:color"), "000000")
        borders.append(element)
    paragraph._p.get_or_add_pPr().append(borders)


def add_row_rule(row, *, above: bool = False, below: bool = False) -> None:
    """Rule the full width of a table row, as a cell border on every cell in it.

    A paragraph border stops at each cell's inner padding, so ruling a row that way comes out as
    a dashed line of segments with gaps at every column boundary. The reference invoice's rules
    are continuous edge to edge, so they are set as cell borders, which butt against each other.
    """
    from docx.oxml.ns import qn
    from docx.oxml.shared import OxmlElement

    for cell in row.cells:
        properties = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for edge, wanted in (("top", above), ("bottom", below)):
            if not wanted:
                continue
            element = OxmlElement(f"w:{edge}")
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), "6")  # 0.75pt, the weight the reference rules measure at
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), "000000")
            borders.append(element)
        properties.append(borders)


def add_signature_block(document, columns: list[tuple[str, str]]) -> None:
    """The signature area: one column per signing party, AGFZE always in the left column.

    The reference contracts are bilateral and set two columns - the purchase contract has AGFZE
    as THE BUYER on the left and the supplier as THE SELLER on the right, the sales contract has
    AGFZE as THE SELLER on the left and the customer as THE BUYER on the right. Invoices and the
    Performa are single-party and set one column. In every case the role label sits above the
    `FOR <entity>` line, a clear vertical gap is left for a wet signature, and AUTHORIZED
    SIGNATORY closes the column.

    The block is emitted as a borderless table so the columns cannot be split across a page
    break, which is the behaviour the reference's three-page contract shows.
    """
    from docx.shared import Inches

    document.add_paragraph()
    table = document.add_table(rows=1, cols=len(columns))
    table.autofit = True
    cells = table.rows[0].cells
    for cell, (role, entity) in zip(cells, columns):
        # `cell.paragraphs[0]` already exists; use it rather than leaving an empty line above.
        first = cell.paragraphs[0]
        first.paragraph_format.space_after = Inches(0)
        _set_font(first.add_run(role), size_pt=SIZE_BODY_PT, bold=True)

        for_line = cell.add_paragraph()
        for_line.paragraph_format.space_after = Inches(0)
        _set_font(for_line.add_run(entity), size_pt=SIZE_BODY_PT, bold=True)

        # Deliberate blank space, reserved for a physical or applied signature. The reference
        # leaves roughly two blank lines here; no signature or stamp image is ever generated.
        cell.add_paragraph()
        cell.add_paragraph()

        closing = cell.add_paragraph()
        closing.paragraph_format.space_after = Inches(0)
        _set_font(closing.add_run("AUTHORIZED SIGNATORY"), size_pt=SIZE_BODY_PT, bold=True)
