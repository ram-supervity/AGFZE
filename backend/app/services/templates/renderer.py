"""Turning a template plus real transaction data into a real DOCX.

Deterministic, start to finish. This module opens a template file that shipped with the platform,
replaces its placeholder runs with values taken from the transaction record, keeps, rewrites or
deletes each clause according to an already-validated plan, and saves the result. There is no
model call in this file and there is no path through it that can reach one: by the time anything
here runs, whatever the AI had to say has already been checked against the template's own clause
registry and reduced to a set of instructions this renderer either understands or refuses.

The only thing it will not do is produce a document with an unresolved placeholder in it. A slot
the transaction has no value for renders as an explicit "not recorded", never as `{{quantity}}`
and never as a plausible-looking blank, because a draft that quietly omits a figure reads exactly
like a draft where that figure is zero.
"""

from __future__ import annotations

import io
import re
from dataclasses import dataclass

from app.services.templates.sales_templates import (
    CLAUSE_MARKER,
    DocumentTemplate,
)

PLACEHOLDER_PATTERN = re.compile(r"\{\{([a-z0-9_]+)\}\}")
MARKER_PATTERN = re.compile(r"\[\[clause:([a-z0-9_]+)\]\]")

# What a slot with nothing behind it says. Explicit, so a reviewer sees the gap.
UNPOPULATED = "- not recorded -"

KEEP = "keep"
REVISE = "revise"
REMOVE = "remove"
ACTIONS = frozenset({KEEP, REVISE, REMOVE})


class TemplateRenderError(Exception):
    """The template could not be rendered, so no document is produced at all."""


@dataclass(frozen=True)
class ClauseDirective:
    key: str
    action: str
    text: str | None = None
    reason: str | None = None


@dataclass(frozen=True)
class RenderResult:
    content: bytes
    kept: list[str]
    revised: list[str]
    removed: list[str]
    # Slots the transaction had no value for. Surfaced so the workspace can tell the reviewer
    # what is missing rather than leaving them to spot it in the document.
    unpopulated: list[str]


def _delete_paragraph(paragraph) -> None:
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def _clause_key(paragraph) -> str | None:
    match = MARKER_PATTERN.search(paragraph.text)
    return match.group(1) if match else None


def _strip_marker(paragraph) -> None:
    for run in paragraph.runs:
        if MARKER_PATTERN.search(run.text):
            run.text = MARKER_PATTERN.sub("", run.text)


def _substitute_run(run, values: dict[str, str], missing: set[str]) -> None:
    """Replace every placeholder inside one run, recording the ones nothing answered.

    The template writes each placeholder as its own run, so this is normally a whole-run swap and
    the run's formatting carries straight through onto the populated value. The regex still
    handles a run that carries surrounding text, because an administrator editing the shipped
    template by hand can merge runs without knowing they have.
    """
    if "{{" not in run.text:
        return

    populated = False

    def replace(match: re.Match[str]) -> str:
        nonlocal populated
        name = match.group(1)
        value = values.get(name)
        if value is None or not str(value).strip():
            missing.add(name)
            return UNPOPULATED
        populated = True
        return str(value)

    run.text = PLACEHOLDER_PATTERN.sub(replace, run.text)
    if populated:
        # Visually distinguishing the deal-specific values from the boilerplate around them, so a
        # reviewer's eye lands on the figures that came out of the transaction.
        run.bold = True
        run.underline = True


def _substitute_paragraph(paragraph, values: dict[str, str], missing: set[str]) -> None:
    for run in paragraph.runs:
        _substitute_run(run, values, missing)


def _rewrite_clause_body(paragraph, text: str, values: dict[str, str], missing: set[str]) -> None:
    """Replace a clause body with revised wording, keeping the template's own formatting.

    The first run is reused so the paragraph keeps the size and face the template gave it; every
    other run goes. Placeholders inside the revised text are then populated exactly as the
    template's own would have been - a revision the model wrote can still refer to the deal's
    real figures, and it still cannot supply them itself.
    """
    runs = list(paragraph.runs)
    if not runs:
        paragraph.add_run(text)
    else:
        runs[0].text = text
        for extra in runs[1:]:
            extra.text = ""
    _substitute_paragraph(paragraph, values, missing)


def render_template(
    template: DocumentTemplate,
    *,
    template_bytes: bytes,
    values: dict[str, str],
    directives: list[ClauseDirective],
) -> RenderResult:
    """Produce the finished DOCX bytes.

    `directives` has already been validated against `template` by the caller - every key is one
    this template declares, no required clause is being removed, and every revision carries real
    text. This function enforces the same invariants a second time anyway, because it is the last
    place before bytes exist and being wrong here means a wrong document rather than a failed job.
    """
    from docx import Document as DocxDocument

    known = template.clause_keys
    required = template.required_clause_keys
    plan: dict[str, ClauseDirective] = {}
    for directive in directives:
        if directive.key not in known:
            raise TemplateRenderError(
                f"'{directive.key}' is not a clause of the {template.document_type} template."
            )
        if directive.action not in ACTIONS:
            raise TemplateRenderError(
                f"'{directive.action}' is not a clause action this renderer performs."
            )
        if directive.action == REMOVE and directive.key in required:
            raise TemplateRenderError(
                f"'{directive.key}' is a required clause of the {template.document_type} "
                "template and cannot be removed."
            )
        if directive.action == REVISE and not (directive.text or "").strip():
            raise TemplateRenderError(
                f"'{directive.key}' was marked for revision with no replacement wording."
            )
        plan[directive.key] = directive

    try:
        document = DocxDocument(io.BytesIO(template_bytes))
    except Exception as exc:  # a corrupt or truncated template file
        raise TemplateRenderError("The document template could not be opened.") from exc

    missing: set[str] = set()
    kept: list[str] = []
    revised: list[str] = []
    removed: list[str] = []

    paragraphs = list(document.paragraphs)
    index = 0
    while index < len(paragraphs):
        paragraph = paragraphs[index]
        key = _clause_key(paragraph)
        if key is None:
            _substitute_paragraph(paragraph, values, missing)
            index += 1
            continue

        # A clause is its heading paragraph and the body paragraph beneath it, exactly as the
        # builder lays it out.
        body = paragraphs[index + 1] if index + 1 < len(paragraphs) else None
        directive = plan.get(key)
        action = directive.action if directive is not None else KEEP

        if action == REMOVE:
            _delete_paragraph(paragraph)
            if body is not None:
                _delete_paragraph(body)
            removed.append(key)
        else:
            _strip_marker(paragraph)
            _substitute_paragraph(paragraph, values, missing)
            if body is not None:
                if action == REVISE and directive is not None and directive.text:
                    _rewrite_clause_body(body, directive.text.strip(), values, missing)
                    revised.append(key)
                else:
                    _substitute_paragraph(body, values, missing)
                    kept.append(key)
            else:
                kept.append(key)

        index += 2 if body is not None else 1

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    _substitute_paragraph(paragraph, values, missing)

    buffer = io.BytesIO()
    document.save(buffer)
    payload = buffer.getvalue()

    # Belt and braces: a produced draft with a live placeholder still in it would be a document
    # that looks finished and is not, which is worse than a failed job.
    if b"{{" in payload and PLACEHOLDER_PATTERN.search(
        "\n".join(p.text for p in DocxDocument(io.BytesIO(payload)).paragraphs)
    ):
        raise TemplateRenderError(
            "The rendered document still contains an unresolved placeholder, so it was discarded."
        )

    return RenderResult(
        content=payload,
        kept=sorted(kept),
        revised=sorted(revised),
        removed=sorted(removed),
        unpopulated=sorted(missing),
    )


def clause_brief(template: DocumentTemplate) -> str:
    """The clause registry as prose, for the prompt. Names every key the model may return."""
    lines = []
    for clause in template.clauses:
        flag = "required, may be revised but never removed" if clause.required else "optional"
        lines.append(f"{clause.key} - {clause.heading} ({flag}). {clause.purpose}")
    return "\n".join(lines)


__all__ = [
    "ACTIONS",
    "CLAUSE_MARKER",
    "KEEP",
    "REMOVE",
    "REVISE",
    "UNPOPULATED",
    "ClauseDirective",
    "RenderResult",
    "TemplateRenderError",
    "clause_brief",
    "render_template",
]
