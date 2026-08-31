"""The generated document templates, as a declaration and as the real DOCX files built from it.

One declaration, two consumers. The builder writes the `.docx` that ships beside this module;
the renderer opens that same `.docx` and populates it. Because the clause registry the AI is
validated against is read from the same declaration, a clause key the model returns that this
file does not name is caught before anything is written - the model cannot invent a clause into
a document, and it cannot delete one the declaration marks as required.

Territory selects field content and the territory-specific reference line, not a separate
template file. **Stated assumption:** the governing material for this step describes what each
territory's paperwork must reference, not four structurally different legal layouts, so building
four separate documents would be inventing three sets of contract structure nobody specified.
One template per document type, with territory-driven content inside it, is what the available
detail actually justifies.
"""

from __future__ import annotations

from dataclasses import dataclass, field
from pathlib import Path

from app.models.enums import DocumentType, Territory
from app.services.templates import house_style

ASSET_ROOT = Path(__file__).resolve().parent / "assets"

# The text width the references lay out against: 8.5in of Letter less the 1.125in left and 1.0in
# right margins the reference `sectPr` declares.
USABLE_WIDTH_IN = (
    house_style.PAGE_WIDTH_IN - house_style.MARGIN_LEFT_IN - house_style.MARGIN_RIGHT_IN
)

# The token the builder writes at the head of a clause paragraph and the renderer strips once it
# has decided what to do with that clause. It is never left in a produced document.
CLAUSE_MARKER = "[[clause:{key}]]"

# What a populated value looks like in the template. Written as one run so substitution is a
# run-level replacement and the surrounding formatting survives it untouched.
PLACEHOLDER = "{{{{{name}}}}}"


@dataclass(frozen=True)
class TemplateField:
    """One deal-specific value the template carries a slot for."""

    name: str
    label: str
    # Fields in the header grid are laid out as a two-column table; body fields are inline.
    in_header: bool = True


@dataclass(frozen=True)
class TemplateClause:
    """One clause of the document, and what the model is allowed to do with it."""

    key: str
    heading: str
    body: str
    # A clause the document is not a document without. The model may revise it; it may never
    # remove it, and a response that asks to is rejected before anything is rendered.
    required: bool = False
    # Told to the model so it can decide. Never rendered.
    purpose: str = ""


@dataclass(frozen=True)
class SignatureColumn:
    """One signing party's column in the signature block.

    `role` is the label the reference puts above the entity - THE SELLER, THE BUYER - and
    `entity_field` names the value that supplies who is signing. AGFZE is always the first
    column in every reference document; the role that AGFZE plays is what changes.
    """

    role: str
    entity_field: str


@dataclass(frozen=True)
class LineItemColumn:
    """One column of a document's line-item table.

    The reference invoices set the column heading over two lines - the name on the first, the
    unit in brackets on the second - so both are carried here. `numeric` drives right alignment,
    which the references apply to quantity, unit price and amount and to nothing else.
    """

    heading: str
    unit: str
    field: str
    numeric: bool = False


@dataclass(frozen=True)
class LineItemTable:
    """The description/quantity/rate/amount table, and the total row beneath it."""

    columns: tuple[LineItemColumn, ...]
    total_label: str
    total_field: str


@dataclass(frozen=True)
class DocumentTemplate:
    document_type: str
    title: str
    subtitle: str
    filename: str
    fields: tuple[TemplateField, ...]
    clauses: tuple[TemplateClause, ...]
    footer_note: str

    # --- layout, measured from the approved reference documents ---------------------------------
    #
    # Every attribute below describes where something sits on the page, and every one of them was
    # read off a reference document rather than chosen. They default to empty so a template that
    # declares none of them still builds - it simply gets the letterhead and nothing else.

    # The `{reference}` ... `DATE: {date}` line the references set under the title and repeat at
    # the head of every continuation page.
    reference_field: str | None = None
    date_field: str | None = None

    # The counterparty's name and address block, as bold flush-left lines.
    party_fields: tuple[str, ...] = ()

    # The line-item table, on the documents that price goods. Contracts do not have one.
    line_items: LineItemTable | None = None

    # Who signs. Two columns on a bilateral contract, one on an invoice.
    signature_columns: tuple[SignatureColumn, ...] = ()

    @property
    def laid_out_field_names(self) -> frozenset[str]:
        """Fields a structural block already places, so the metadata block does not repeat them."""
        placed = set(self.party_fields)
        for name in (self.reference_field, self.date_field):
            if name:
                placed.add(name)
        if self.line_items is not None:
            placed.update(column.field for column in self.line_items.columns)
            placed.add(self.line_items.total_field)
        return frozenset(placed)

    @property
    def clause_keys(self) -> frozenset[str]:
        return frozenset(clause.key for clause in self.clauses)

    @property
    def required_clause_keys(self) -> frozenset[str]:
        return frozenset(clause.key for clause in self.clauses if clause.required)

    def clause(self, key: str) -> TemplateClause | None:
        return next((row for row in self.clauses if row.key == key), None)

    @property
    def field_names(self) -> tuple[str, ...]:
        return tuple(row.name for row in self.fields)


# --- territory-driven content -----------------------------------------------------------------
#
# Content, not layout. Each entry is the reference line the destination territory's customs and
# import regime requires the paperwork to carry, and it is the only thing that changes between a
# China-bound and an India-bound draft of the same document.

TERRITORY_REFERENCES: dict[str, str] = {
    Territory.INDIA.value: (
        "Consignment is declared for import into India. Buyer's IEC and GSTIN are to be quoted "
        "on all shipping documents, and the pre-shipment inspection certificate required under "
        "the Hazardous and Other Wastes (Management and Transboundary Movement) Rules is to "
        "accompany this consignment."
    ),
    Territory.CHINA.value: (
        "Consignment is declared for import into the People's Republic of China. Goods are "
        "described in accordance with GB/T 38470 recycled non-ferrous metal raw material "
        "standards, and Buyer's China Customs registration and AQSIQ / GACC filing details are "
        "to be quoted on all shipping documents."
    ),
    Territory.JAPAN.value: (
        "Consignment is declared for import into Japan. Buyer is to quote its Japan Customs "
        "importer code, and the goods description is to follow the Japanese Customs tariff "
        "classification agreed between the parties."
    ),
    Territory.OTHER.value: (
        "Consignment is declared for import into the destination named above. Buyer is "
        "responsible for all import licences, registrations and customs formalities required at "
        "the port of discharge, and is to advise Seller of any documentary requirement in "
        "writing before shipment."
    ),
}


def territory_reference(territory: str | None) -> str:
    """The destination's required reference wording. Never a guess - `other` is a real entry."""
    return TERRITORY_REFERENCES.get(
        (territory or "").strip().lower(), TERRITORY_REFERENCES[Territory.OTHER.value]
    )


# --- the sales contract -------------------------------------------------------------------------

CONTRACT_FIELDS: tuple[TemplateField, ...] = (
    TemplateField("contract_no", "Sales contract number"),
    TemplateField("contract_date", "Date"),
    TemplateField("batch_number", "AGFZE batch reference"),
    TemplateField("seller", "Seller"),
    TemplateField("buyer", "Buyer"),
    TemplateField("territory", "Destination territory"),
    TemplateField("commodity", "Commodity"),
    TemplateField("commodity_code", "Trade grade"),
    TemplateField("quantity", "Quantity (this shipment)"),
    TemplateField("contracted_quantity", "Contracted quantity (total)"),
    TemplateField("price_terms", "Price"),
    TemplateField("currency", "Currency"),
    TemplateField("payment_condition", "Payment condition"),
    TemplateField("port_of_loading", "Port of loading"),
    TemplateField("port_of_discharge", "Port of discharge"),
    TemplateField("inland_container_depot", "Inland container depot"),
    TemplateField("bl_reference", "Bill of lading reference"),
)

CONTRACT_CLAUSES: tuple[TemplateClause, ...] = (
    TemplateClause(
        key="parties",
        heading="1. Parties and subject",
        body=(
            "This contract is made between {{seller}} (“Seller”) and {{buyer}} "
            "(“Buyer”) for the sale and purchase of the goods described below, shipped "
            "under AGFZE batch reference {{batch_number}}."
        ),
        required=True,
        purpose="Names the two parties and the batch. Always present in any sales contract.",
    ),
    TemplateClause(
        key="goods",
        heading="2. Goods",
        body=(
            "Commodity: {{commodity}} (trade grade {{commodity_code}}). Quantity for this "
            "shipment: {{quantity}}, forming part of a total contracted quantity of "
            "{{contracted_quantity}} under this contract number. Material is sold on an as-is "
            "recovered-metal basis, free of radioactivity, explosives and closed containers."
        ),
        required=True,
        purpose="Describes what is being sold and how much. Always present.",
    ),
    TemplateClause(
        key="quantity_tolerance",
        heading="3. Quantity tolerance",
        body=(
            "Shipped weight may vary from the quantity stated above within the tolerance agreed "
            "between the parties. The final settled weight is the outturn weight established at "
            "the port of discharge unless the parties have agreed otherwise in writing."
        ),
        purpose=(
            "Applies where the shipped weight is expected to settle against an outturn weight. "
            "Remove it where the parties have contracted a fixed, non-varying quantity."
        ),
    ),
    TemplateClause(
        key="pricing_fixed",
        heading="4. Price",
        body=(
            "The price for this shipment is {{price_terms}}, on the delivery terms stated "
            "above. No further price fixation applies to this contract."
        ),
        purpose=(
            "The fixed-price clause. Correct where the deal is priced at a flat rate, and also "
            "where an LME-linked deal has since been fixed by the customer, in which case the "
            "fixed rate and the fixation date should be stated."
        ),
    ),
    TemplateClause(
        key="pricing_lme",
        heading="4. Price",
        body=(
            "The price for this shipment is {{price_terms}}, calculated on the London Metal "
            "Exchange cash settlement for the relevant grade over the pricing period agreed "
            "between the parties, settled in {{currency}}, on the delivery terms stated above."
        ),
        purpose=(
            "The LME-percentage clause. Correct only while the price is still a percentage of "
            "the exchange and the customer has not fixed. Remove it once the customer has fixed."
        ),
    ),
    TemplateClause(
        key="price_fixation",
        heading="5. Price fixation",
        body=(
            "Buyer shall declare its fixation in writing to Seller during the agreed pricing "
            "period. Where Buyer has not declared a fixation by the close of that period, the "
            "price shall be established on the basis agreed between the parties."
        ),
        purpose=(
            "Explains how the customer fixes a price that is not yet fixed. Remove it entirely "
            "once fixation has already happened, because there is nothing left to fix."
        ),
    ),
    TemplateClause(
        key="payment_cad",
        heading="6. Payment",
        body=(
            "Payment is on Cash Against Documents ({{payment_condition}}). Seller shall present "
            "the full documentary set through the agreed channel, and Buyer shall pay against "
            "presentation of those documents."
        ),
        purpose=(
            "The Cash Against Documents payment clause. Use it when the payment condition is "
            "CAD; remove it when the condition is TT."
        ),
    ),
    TemplateClause(
        key="payment_tt",
        heading="6. Payment",
        body=(
            "Payment is by telegraphic transfer ({{payment_condition}}) to Seller's nominated "
            "account, in {{currency}} and in cleared funds, on the terms agreed between the "
            "parties. Release of the transport documents follows receipt of cleared funds."
        ),
        purpose=(
            "The telegraphic-transfer payment clause. Use it when the payment condition is TT; "
            "remove it when the condition is CAD."
        ),
    ),
    TemplateClause(
        key="delivery",
        heading="7. Delivery and shipment",
        body=(
            "Shipment is from {{port_of_loading}} to {{port_of_discharge}}. Where an inland "
            "container depot is nominated, it is {{inland_container_depot}}. Bill of lading "
            "reference for this shipment: {{bl_reference}}."
        ),
        required=True,
        purpose="States where the cargo moves from and to. Always present.",
    ),
    TemplateClause(
        key="documents",
        heading="8. Documents and destination requirements",
        body=(
            "Seller shall provide the commercial invoice, packing list, bill of lading and "
            "certificate of origin, together with any further document the destination requires. "
            "{{territory_reference}}"
        ),
        required=True,
        purpose=(
            "The documentary set and the destination's own requirements. Always present; its "
            "destination wording is populated from the recorded territory rather than chosen."
        ),
    ),
    TemplateClause(
        key="inspection",
        heading="9. Inspection",
        body=(
            "Either party may appoint an independent surveyor at the port of loading or the port "
            "of discharge, at its own cost, and the other party shall be given reasonable notice "
            "of any such appointment."
        ),
        purpose=(
            "The surveyor clause. Keep it where an independent inspection is part of the deal; "
            "remove it where the parties have agreed to settle on the carrier's figures alone."
        ),
    ),
    TemplateClause(
        key="title_risk",
        heading="10. Title and risk",
        body=(
            "Risk passes in accordance with the delivery terms stated above. Title passes to "
            "Buyer on receipt by Seller of payment in full and in cleared funds."
        ),
        required=True,
        purpose="When risk and ownership move. Always present.",
    ),
    TemplateClause(
        key="force_majeure",
        heading="11. Force majeure",
        body=(
            "Neither party is liable for a failure to perform caused by an event beyond its "
            "reasonable control. The affected party shall notify the other in writing without "
            "delay and the parties shall discuss the consequences in good faith."
        ),
        required=True,
        purpose="Standard force majeure. Always present.",
    ),
    TemplateClause(
        key="governing_law",
        heading="12. Governing law",
        body=(
            "This contract is governed by the law agreed between the parties, and any dispute "
            "arising out of it is to be resolved in the forum the parties have agreed."
        ),
        required=True,
        purpose="Governing law and forum. Always present.",
    ),
)

SALES_CONTRACT_TEMPLATE = DocumentTemplate(
    document_type=DocumentType.DRAFT_CONTRACT.value,
    title="SALES CONTRACT",
    subtitle="Draft for internal review - not issued, not signed",
    filename="sales_contract_template.docx",
    fields=CONTRACT_FIELDS,
    clauses=CONTRACT_CLAUSES,
    # Layout from `Sales Contract_2026-27 COPY 1.xlsx`, sheets INDIA (QP) / CHINA (QP): contract
    # number left and DATE right under the title, the buyer's name and address beneath, and a
    # two-column signature block with AGFZE as THE SELLER on the left.
    reference_field="contract_no",
    date_field="contract_date",
    party_fields=("buyer",),
    signature_columns=(
        SignatureColumn("THE SELLER", "agfze"),
        SignatureColumn("THE BUYER", "buyer"),
    ),
    footer_note=(
        "DRAFT. Generated by the AGFZE Command Centre from the transaction record on "
        "{{generated_at}} at the request of {{generated_by}}. This document has not been "
        "issued, sent or signed. It is for review inside AGFZE only; the signed original is a "
        "wet-signed paper document produced outside this platform."
    ),
)


# --- the sales invoice ---------------------------------------------------------------------------

INVOICE_FIELDS: tuple[TemplateField, ...] = (
    TemplateField("invoice_no", "Sales invoice number"),
    TemplateField("invoice_date", "Date"),
    TemplateField("contract_no", "Sales contract number"),
    TemplateField("batch_number", "AGFZE batch reference"),
    TemplateField("seller", "Seller"),
    TemplateField("buyer", "Buyer"),
    TemplateField("territory", "Destination territory"),
    TemplateField("commodity", "Commodity"),
    TemplateField("commodity_code", "Trade grade"),
    TemplateField("quantity", "Quantity"),
    TemplateField("price_terms", "Rate"),
    TemplateField("currency", "Currency"),
    TemplateField("total_value", "Invoice value"),
    TemplateField("invoice_basis", "Invoice basis"),
    TemplateField("payment_condition", "Payment condition"),
    TemplateField("port_of_loading", "Port of loading"),
    TemplateField("port_of_discharge", "Port of discharge"),
    TemplateField("bl_reference", "Bill of lading reference"),
)

INVOICE_CLAUSES: tuple[TemplateClause, ...] = (
    TemplateClause(
        key="header",
        heading="Invoice to",
        body=(
            "{{buyer}}, against sales contract {{contract_no}} and AGFZE batch reference "
            "{{batch_number}}. Bill of lading reference {{bl_reference}}."
        ),
        required=True,
        purpose="Who is invoiced and against what. Always present.",
    ),
    TemplateClause(
        key="goods_description",
        heading="Description of goods",
        body=(
            "{{commodity}} (trade grade {{commodity_code}}), {{quantity}}, shipped "
            "{{port_of_loading}} to {{port_of_discharge}}."
        ),
        required=True,
        purpose="What is being invoiced. Always present.",
    ),
    TemplateClause(
        key="pricing_provisional",
        heading="Provisional value",
        body=(
            "This is a PROVISIONAL invoice. The value of {{total_value}} is calculated on a "
            "price of {{price_terms}} and is subject to adjustment once the price is fixed. A "
            "final invoice will follow on fixation."
        ),
        purpose=(
            "The provisional-value clause. Correct while the customer has not fixed the price. "
            "Remove it once fixation has happened, because the value is then settled."
        ),
    ),
    TemplateClause(
        key="pricing_final",
        heading="Final value",
        body=(
            "This is a FINAL invoice. The value of {{total_value}} is calculated on a fixed "
            "price of {{price_terms}} and is not subject to further adjustment."
        ),
        purpose=(
            "The final-value clause. Correct once the customer has fixed the price. Remove it "
            "while the price is still unfixed."
        ),
    ),
    TemplateClause(
        key="lme_reference",
        heading="Pricing basis",
        body=(
            "Price basis: {{price_terms}}, referenced to the London Metal Exchange cash "
            "settlement for the relevant grade over the agreed pricing period."
        ),
        purpose=(
            "The exchange-reference line. Keep it on any LME-linked deal, fixed or not, since "
            "the customer still needs the basis on the face of the invoice. Remove it entirely "
            "on a flat-priced deal, where there is no exchange reference to state."
        ),
    ),
    TemplateClause(
        key="payment_cad",
        heading="Payment",
        body=(
            "Payable on Cash Against Documents ({{payment_condition}}) in {{currency}}, against "
            "presentation of the full documentary set."
        ),
        purpose="Use where the payment condition is CAD; remove where it is TT.",
    ),
    TemplateClause(
        key="payment_tt",
        heading="Payment",
        body=(
            "Payable by telegraphic transfer ({{payment_condition}}) in {{currency}} to the "
            "Seller's nominated account, in cleared funds."
        ),
        purpose="Use where the payment condition is TT; remove where it is CAD.",
    ),
    TemplateClause(
        key="bank_details",
        heading="Remittance",
        body=(
            "Bank details for remittance are those held on file for {{seller}} and confirmed to "
            "Buyer in writing. Buyer must not act on bank details received by any other route."
        ),
        required=True,
        purpose=(
            "Remittance instructions and the warning against acting on details from elsewhere. "
            "Always present; it is a fraud control, not a commercial term."
        ),
    ),
    TemplateClause(
        key="destination_declaration",
        heading="Destination requirements",
        body="{{territory_reference}}",
        required=True,
        purpose=(
            "The destination's own declaration. Always present; its wording is populated from "
            "the recorded territory rather than chosen."
        ),
    ),
    TemplateClause(
        key="declaration",
        heading="Declaration",
        body=(
            "We certify the above particulars to be true and correct, and that the goods are of "
            "the origin and description stated."
        ),
        required=True,
        purpose="The seller's certification. Always present.",
    ),
)

SALES_INVOICE_TEMPLATE = DocumentTemplate(
    document_type=DocumentType.DRAFT_INVOICE.value,
    title="COMMERCIAL INVOICE",
    subtitle="Draft for internal review - not issued, not signed",
    filename="sales_invoice_template.docx",
    fields=INVOICE_FIELDS,
    clauses=INVOICE_CLAUSES,
    # Layout from page 1 of the reference invoice PDF, the COMMERCIAL INVOICE. Column headings,
    # their bracketed units and the total label are that document's own wording; the currency is
    # a placeholder because the reference states USD and the platform trades in more than one.
    reference_field="invoice_no",
    date_field="invoice_date",
    party_fields=("buyer",),
    line_items=LineItemTable(
        columns=(
            LineItemColumn("DESCRIPTION", "", "commodity"),
            LineItemColumn("QUANTITY", "(NET MTS)", "quantity", numeric=True),
            LineItemColumn("UNIT PRICE", "({{currency}} PER MT)", "unit_rate", numeric=True),
            LineItemColumn("AMOUNT", "({{currency}})", "total_value", numeric=True),
        ),
        total_label="{{currency}} TOTAL PAYABLE",
        total_field="total_value",
    ),
    signature_columns=(SignatureColumn("", "agfze"),),
    footer_note=(
        "DRAFT. Generated by the AGFZE Command Centre from the transaction record on "
        "{{generated_at}} at the request of {{generated_by}}. This document has not been "
        "issued, sent or signed. It is for review inside AGFZE only; the signed original is a "
        "wet-signed paper document produced outside this platform."
    ),
)

# --- the Performa invoice ---------------------------------------------------------------------
#
# The advance, clean invoice, raised before the cargo is weighed. Discovery describes it exactly
# that way, and its defining property is what it *lacks*: there is no weight slip behind it and no
# shipped weight to state, which is why it is a distinct document rather than a commercial invoice
# with some fields left blank. A commercial invoice with a blank weight would read as a document
# somebody forgot to finish; this one is complete as it stands.
#
# It therefore carries no `bl_reference` and no shipped quantity: the contracted quantity is what
# it invoices against. Every clause below that would have asserted something about the shipment
# has been left out rather than softened.
#
# **Approval tier - an open question, stated here rather than decided.** Discovery says a Performa
# invoice "requires CEO approval". This platform has no `ceo` role: its approving tier is
# `approver_hod`, with `admin` alongside. Inventing a role would change who can sign off what on
# the day the platform goes live, which is not a decision a template file should make. So a
# Performa invoice routes through the existing `approver_hod` tier, exactly as every other draft
# does, and the question is recorded in docs/KNOWN-GAPS.md for AGFZE to answer.

PERFORMA_FIELDS: tuple[TemplateField, ...] = (
    TemplateField("invoice_no", "Performa invoice number"),
    TemplateField("invoice_date", "Date"),
    TemplateField("contract_no", "Sales contract number"),
    TemplateField("batch_number", "AGFZE batch reference"),
    TemplateField("seller", "Seller"),
    TemplateField("buyer", "Buyer"),
    TemplateField("territory", "Destination territory"),
    TemplateField("commodity", "Commodity"),
    TemplateField("commodity_code", "Trade grade"),
    TemplateField("quantity", "Contracted quantity"),
    TemplateField("price_terms", "Rate"),
    TemplateField("currency", "Currency"),
    TemplateField("total_value", "Performa value"),
    TemplateField("payment_condition", "Payment condition"),
    TemplateField("port_of_loading", "Port of loading"),
    TemplateField("port_of_discharge", "Port of discharge"),
)

PERFORMA_CLAUSES: tuple[TemplateClause, ...] = (
    TemplateClause(
        key="header",
        heading="Performa invoice to",
        body=(
            "{{buyer}}, against sales contract {{contract_no}} and AGFZE batch reference "
            "{{batch_number}}."
        ),
        required=True,
        purpose="Who is invoiced and against what. Always present.",
    ),
    TemplateClause(
        key="provisional_basis",
        heading="Basis",
        body=(
            "This is a Performa invoice raised in advance of shipment. The quantity stated is "
            "the contracted quantity; it is not a shipped weight and no weight slip accompanies "
            "this document. A commercial invoice stating the shipped weight follows once the "
            "cargo is loaded and weighed."
        ),
        required=True,
        purpose=(
            "The clause that makes this document honest about what it is. Always present, and "
            "never to be revised into a claim about a shipped weight - there is not one."
        ),
    ),
    TemplateClause(
        key="goods",
        heading="Goods",
        body=(
            "{{commodity}} ({{commodity_code}}), {{quantity}} contracted, at {{price_terms}}, "
            "{{currency}} {{total_value}}."
        ),
        required=True,
        purpose="The commercial particulars as contracted. Always present.",
    ),
    TemplateClause(
        key="advance_payment",
        heading="Advance payment",
        body=(
            "Payable in advance in {{currency}} on the terms agreed ({{payment_condition}}), to "
            "the Seller's nominated account in cleared funds."
        ),
        required=True,
        purpose=(
            "A Performa invoice exists to be paid against in advance, so the payment clause is "
            "always present. Revise the wording to match the agreed terms; never remove it."
        ),
    ),
    TemplateClause(
        key="bank_details",
        heading="Remittance",
        body=(
            "Bank details for remittance are those held on file for {{seller}} and confirmed to "
            "Buyer in writing. Buyer must not act on bank details received by any other route."
        ),
        required=True,
        purpose=(
            "Remittance instructions and the warning against acting on details from elsewhere. "
            "Always present; it is a fraud control, not a commercial term."
        ),
    ),
    TemplateClause(
        key="shipment",
        heading="Shipment",
        body="Shipment from {{port_of_loading}} to {{port_of_discharge}}.",
        purpose=(
            "The intended routing. Keep it where both ports are recorded; remove it where the "
            "routing is not yet agreed, rather than stating a port nobody has confirmed."
        ),
    ),
    TemplateClause(
        key="destination_declaration",
        heading="Destination requirements",
        body="{{territory_reference}}",
        required=True,
        purpose=(
            "The destination's own declaration. Always present; its wording is populated from "
            "the recorded territory rather than chosen."
        ),
    ),
)

PERFORMA_INVOICE_TEMPLATE = DocumentTemplate(
    document_type=DocumentType.DRAFT_PERFORMA_INVOICE.value,
    title="PERFORMA INVOICE",
    subtitle="Draft for internal review - not issued, not signed",
    filename="performa_invoice_template.docx",
    fields=PERFORMA_FIELDS,
    clauses=PERFORMA_CLAUSES,
    # Layout from the PERFORMA INVOICE sheet of the reference contract workbook. It heads itself
    # with the *contract* number rather than an invoice number - the reference does the same,
    # because a Performa is raised against the contract before an invoice number exists - and
    # signs single-party.
    #
    # The reference's total row is `ADVANCE 10%`, an advance percentage of the line value. This
    # platform records no advance percentage anywhere, so the row states the full contracted
    # value instead. Recorded as a data gap rather than reproduced with an invented 10%.
    reference_field="contract_no",
    date_field="invoice_date",
    party_fields=("buyer",),
    line_items=LineItemTable(
        columns=(
            LineItemColumn("ITEM DESCRIPTION", "", "commodity"),
            LineItemColumn("QTY", "(MT)", "quantity", numeric=True),
            LineItemColumn("PRICE", "({{currency}})", "unit_rate", numeric=True),
            LineItemColumn("AMOUNT", "({{currency}})", "total_value", numeric=True),
        ),
        total_label="{{currency}} TOTAL",
        total_field="total_value",
    ),
    signature_columns=(SignatureColumn("", "agfze"),),
    footer_note=(
        "DRAFT. Generated by the AGFZE Command Centre from the transaction record on "
        "{{generated_at}} at the request of {{generated_by}}. This document has not been "
        "issued, sent or signed. It is a Performa invoice raised in advance of shipment and "
        "states a contracted quantity, not a shipped weight."
    ),
)


# --- the bank cover letter ----------------------------------------------------------------------
#
# The covering note that goes to the bank with a documentary set. Discovery describes it as
# auto-filled, which is accurate: every value on it is already recorded on the transaction, and it
# asserts nothing that is not stated on the documents it accompanies.
#
# Deliberately narrow. It states what is enclosed, against which contract, and what the bank is
# asked to do with it. It carries **no** commercial terms of its own - no rate, no invoice value -
# because a cover letter that restated the commercial terms would be a second, unsigned source of
# them, and a discrepancy between it and the invoice is exactly the sort of thing that holds a
# documentary presentation up at the bank.

BANK_LETTER_FIELDS: tuple[TemplateField, ...] = (
    TemplateField("letter_date", "Date"),
    TemplateField("contract_no", "Sales contract number"),
    TemplateField("invoice_no", "Sales invoice number"),
    TemplateField("batch_number", "AGFZE batch reference"),
    TemplateField("seller", "Seller"),
    TemplateField("buyer", "Buyer"),
    TemplateField("bank_name", "Presenting bank"),
    TemplateField("commodity", "Commodity"),
    TemplateField("quantity", "Quantity"),
    TemplateField("currency", "Currency"),
    TemplateField("total_value", "Invoice value"),
    TemplateField("payment_condition", "Payment condition"),
    TemplateField("bl_reference", "Bill of lading reference"),
    TemplateField("port_of_discharge", "Port of discharge"),
)

BANK_LETTER_CLAUSES: tuple[TemplateClause, ...] = (
    TemplateClause(
        key="addressee",
        heading="To",
        body=(
            "{{bank_name}}. We enclose the documentary set covering our sales contract "
            "{{contract_no}} and AGFZE batch reference {{batch_number}}, drawn on {{buyer}}."
        ),
        required=True,
        purpose="Who the letter is to and what it covers. Always present.",
    ),
    TemplateClause(
        key="enclosures",
        heading="Documents enclosed",
        body=(
            "Commercial invoice {{invoice_no}}, bill of lading {{bl_reference}}, and the "
            "supporting certificates for {{quantity}} of {{commodity}} shipped to "
            "{{port_of_discharge}}."
        ),
        required=True,
        purpose=(
            "The list of what is actually in the envelope. Always present, and revise it to "
            "match the set genuinely being presented - a cover letter that lists a document "
            "which is not enclosed is what holds a presentation up."
        ),
    ),
    TemplateClause(
        key="instruction",
        heading="Instruction",
        body=(
            "Please present these documents on {{payment_condition}} terms and release them "
            "against payment of {{currency}} {{total_value}} in accordance with the contract."
        ),
        required=True,
        purpose=(
            "What the bank is being asked to do. Always present. The one place a value appears "
            "on this letter, because the release is conditional on it."
        ),
    ),
    TemplateClause(
        key="remittance",
        heading="Remittance",
        body=(
            "Proceeds are to be remitted to the account held on file for {{seller}}. Please "
            "confirm receipt of these documents and advise us on presentation."
        ),
        required=True,
        purpose="Where the money goes and the acknowledgement asked for. Always present.",
    ),
)

BANK_COVER_LETTER_TEMPLATE = DocumentTemplate(
    document_type=DocumentType.DRAFT_BANK_COVER_LETTER.value,
    title="BANK COVER LETTER",
    subtitle="Draft for internal review - not issued, not signed",
    filename="bank_cover_letter_template.docx",
    fields=BANK_LETTER_FIELDS,
    clauses=BANK_LETTER_CLAUSES,
    # **No reference document.** The supplied reference set contains no bank cover letter, so the
    # layout below is the house style the other four references establish - letterhead, reference
    # line, addressee block, single signature - and not a replica of anything. It should be
    # checked against a real one before it is relied on.
    reference_field="contract_no",
    date_field="letter_date",
    party_fields=("bank_name",),
    signature_columns=(SignatureColumn("", "agfze"),),
    footer_note=(
        "DRAFT. Generated by the AGFZE Command Centre from the transaction record on "
        "{{generated_at}} at the request of {{generated_by}}. This document has not been "
        "issued, sent or signed, and has not been presented to any bank."
    ),
)

PURCHASE_CONTRACT_FIELDS: tuple[TemplateField, ...] = (
    TemplateField("contract_no", "Purchase contract number"),
    TemplateField("contract_date", "Date"),
    TemplateField("batch_number", "AGFZE batch reference"),
    TemplateField("buyer", "Buyer"),
    TemplateField("seller", "Seller / Supplier"),
    TemplateField("commodity", "Commodity"),
    TemplateField("commodity_code", "Trade grade"),
    TemplateField("quantity", "Quantity"),
    TemplateField("price_terms", "Price"),
    TemplateField("currency", "Currency"),
    TemplateField("total_value", "Total purchase value"),
    TemplateField("payment_condition", "Payment condition"),
    TemplateField("port_of_loading", "Port of loading"),
    TemplateField("invoice_basis", "Invoice basis"),
)

PURCHASE_CONTRACT_CLAUSES: tuple[TemplateClause, ...] = (
    TemplateClause(
        key="parties",
        heading="1. Parties and subject",
        body=(
            "This purchase contract is made between {{buyer}} (“Buyer”) and {{seller}} "
            "(“Seller”) for the purchase and sale of the goods described below, under "
            "AGFZE batch reference {{batch_number}}."
        ),
        required=True,
        purpose="Names Buyer (AGFZE) and Seller/Supplier and the batch. Always present.",
    ),
    TemplateClause(
        key="goods",
        heading="2. Goods and specifications",
        body=(
            "Commodity: {{commodity}} (trade grade {{commodity_code}}). Quantity: {{quantity}}. "
            "Material is supplied on an as-is recovered metal basis, free from radioactive materials, "
            "explosives, arms, ammunition, and deleterious substances."
        ),
        required=True,
        purpose="Describes what is being bought and how much. Always present.",
    ),
    TemplateClause(
        key="pricing",
        heading="3. Price and value",
        body=(
            "The agreed price is {{price_terms}}, with an estimated total value of "
            "{{total_value}}, settled in {{currency}} on {{invoice_basis}} basis."
        ),
        required=True,
        purpose="States the price, total amount and currency. Always present.",
    ),
    TemplateClause(
        key="delivery",
        heading="4. Shipment and loading",
        body=(
            "Shipment is from {{port_of_loading}}. Supplier shall arrange prompt loading and "
            "issue commercial shipping documents upon vessel departure."
        ),
        required=True,
        purpose="States shipping port and loading terms. Always present.",
    ),
    TemplateClause(
        key="payment",
        heading="5. Payment terms",
        body=(
            "Payment shall be made on {{payment_condition}} terms in {{currency}} against "
            "presentation of clean transport documents, commercial invoice and packing list."
        ),
        required=True,
        purpose="Payment conditions and documentation. Always present.",
    ),
    TemplateClause(
        key="inspection",
        heading="6. Weight and quality determination",
        body=(
            "Settlement of final weight and quality shall be determined on destination outturn "
            "weights and independent inspection report unless agreed otherwise in writing."
        ),
        purpose="Outturn weight and inspection terms.",
    ),
    TemplateClause(
        key="title_risk",
        heading="7. Title and risk",
        body=(
            "Risk passes in accordance with agreed delivery terms. Title passes to Buyer upon "
            "release of documents or receipt of payment in cleared funds."
        ),
        required=True,
        purpose="Title and risk transfer. Always present.",
    ),
    TemplateClause(
        key="governing_law",
        heading="8. Governing law",
        body=(
            "This contract is governed by the laws agreed between the parties, and any dispute "
            "shall be resolved in the forum agreed between the parties."
        ),
        required=True,
        purpose="Governing law and dispute resolution. Always present.",
    ),
)

PURCHASE_CONTRACT_TEMPLATE = DocumentTemplate(
    document_type=DocumentType.DRAFT_PURCHASE_CONTRACT.value,
    title="PURCHASE CONTRACT",
    subtitle="Draft for internal review - not issued, not signed",
    filename="purchase_contract_template.docx",
    fields=PURCHASE_CONTRACT_FIELDS,
    clauses=PURCHASE_CONTRACT_CLAUSES,
    # Layout from `Sample Purchase Contract 1.docx`. The signature roles are the mirror of the
    # sales contract's: AGFZE is THE BUYER here, and it keeps the left column, which is what the
    # reference's third page shows.
    reference_field="contract_no",
    date_field="contract_date",
    party_fields=("supplier",),
    signature_columns=(
        SignatureColumn("THE BUYER", "agfze"),
        SignatureColumn("THE SELLER", "supplier"),
    ),
    footer_note=(
        "DRAFT. Generated by the AGFZE Command Centre from the purchase transaction record on "
        "{{generated_at}} at the request of {{generated_by}}. This document has not been "
        "issued, sent or signed. It is for review inside AGFZE only; the signed original is a "
        "wet-signed paper document produced outside this platform."
    ),
)


# --- the cost sheet -----------------------------------------------------------------------------

COST_SHEET_FIELDS: tuple[TemplateField, ...] = (
    TemplateField("batch_number", "AGFZE batch reference"),
    TemplateField("date", "Date"),
    TemplateField("supplier", "Supplier"),
    TemplateField("commodity", "Commodity"),
    TemplateField("commodity_code", "Trade grade"),
    TemplateField("quantity", "Quantity"),
    TemplateField("purchase_rate", "Purchase rate"),
    TemplateField("purchase_value", "Purchase cost"),
    TemplateField("freight_and_logistics", "Logistics allowance"),
    TemplateField("financing_and_charges", "Finance & bank charges"),
    TemplateField("total_cost", "Total landed cost"),
    TemplateField("currency", "Currency"),
    TemplateField("invoice_basis", "Invoice basis"),
    TemplateField("port_of_loading", "Port of loading"),
)

COST_SHEET_CLAUSES: tuple[TemplateClause, ...] = (
    TemplateClause(
        key="deal_summary",
        heading="1. Transaction and cargo summary",
        body=(
            "Cost estimation for batch {{batch_number}}. Commodity: {{commodity}} ({{commodity_code}}), "
            "quantity: {{quantity}}. Supplier: {{supplier}}, loading port: {{port_of_loading}}."
        ),
        required=True,
        purpose="Deal header and cargo particulars. Always present.",
    ),
    TemplateClause(
        key="procurement_cost",
        heading="2. Purchase procurement cost",
        body=(
            "Procurement rate: {{purchase_rate}}. Total purchase cost: {{purchase_value}} "
            "({{invoice_basis}} basis, settled in {{currency}})."
        ),
        required=True,
        purpose="Base procurement rate and purchase value. Always present.",
    ),
    TemplateClause(
        key="freight_logistics",
        heading="3. Freight and handling provision",
        body=(
            "Ocean freight, container handling, clearance and local port logistics provision: "
            "{{freight_and_logistics}}."
        ),
        required=True,
        purpose="Logistics and freight charges provision. Always present.",
    ),
    TemplateClause(
        key="financing_hedging",
        heading="4. Hedging and financial charges",
        body=(
            "Hedging and price basis: {{hedge_details}}. Banking presentation, document release "
            "and financing costs allowance: {{financing_and_charges}}."
        ),
        required=True,
        purpose="Financial and hedging cost analysis. Always present.",
    ),
    TemplateClause(
        key="landed_cost_summary",
        heading="5. Projected landed cost and structure",
        body=(
            "Total projected cost for batch {{batch_number}} is {{total_cost}} {{currency}}. "
            "{{b2b_split}}"
        ),
        required=True,
        purpose="Total landed cost summary and partnership structure. Always present.",
    ),
)

COST_SHEET_TEMPLATE = DocumentTemplate(
    document_type=DocumentType.DRAFT_COST_SHEET.value,
    title="TRANSACTION COST SHEET",
    subtitle="Internal Review Only - Cost & Margin Analysis",
    filename="cost_sheet_template.docx",
    fields=COST_SHEET_FIELDS,
    clauses=COST_SHEET_CLAUSES,
    # The reference cost sheet (`Sample Cost Sheet 1.xlsx`) is an unbranded landscape spreadsheet
    # grid with a payments ledger beside the cost build-up - a working document, not a letterhead
    # one. This template is a portrait letterhead document and so does **not** replicate it; see
    # the replication report for why, and for what replicating it would actually take. What is
    # reproduced here is the reference's field set and its internal-only character.
    reference_field="batch_number",
    date_field="date",
    party_fields=("supplier",),
    footer_note=(
        "INTERNAL COST SHEET. Generated by the AGFZE Command Centre from the purchase transaction "
        "record on {{generated_at}} at the request of {{generated_by}}. For internal calculation and "
        "financial review only. Confidential."
    ),
)


TEMPLATES_BY_TYPE: dict[str, DocumentTemplate] = {
    SALES_CONTRACT_TEMPLATE.document_type: SALES_CONTRACT_TEMPLATE,
    SALES_INVOICE_TEMPLATE.document_type: SALES_INVOICE_TEMPLATE,
    PERFORMA_INVOICE_TEMPLATE.document_type: PERFORMA_INVOICE_TEMPLATE,
    BANK_COVER_LETTER_TEMPLATE.document_type: BANK_COVER_LETTER_TEMPLATE,
    PURCHASE_CONTRACT_TEMPLATE.document_type: PURCHASE_CONTRACT_TEMPLATE,
    COST_SHEET_TEMPLATE.document_type: COST_SHEET_TEMPLATE,
}

# Fields every template carries whether or not it declares them: the provenance line at the foot
# of a generated draft, and the destination wording clauses reference.
COMMON_FIELDS: tuple[str, ...] = ("generated_at", "generated_by", "territory_reference")


@dataclass(frozen=True)
class BuildReport:
    written: list[str] = field(default_factory=list)
    skipped: list[str] = field(default_factory=list)


def template_path(template: DocumentTemplate) -> Path:
    return ASSET_ROOT / template.filename


def _split_placeholders(text: str) -> list[tuple[str, bool]]:
    """Cut a body into alternating literal and placeholder segments.

    Each placeholder becomes its own run in the built template, which is the whole reason
    substitution at render time is a run-level replacement rather than a fragile scan across a
    paragraph whose text `python-docx` has already split at formatting boundaries.
    """
    segments: list[tuple[str, bool]] = []
    remainder = text
    while True:
        start = remainder.find("{{")
        if start == -1:
            break
        end = remainder.find("}}", start)
        if end == -1:
            break
        if start:
            segments.append((remainder[:start], False))
        segments.append((remainder[start : end + 2], True))
        remainder = remainder[end + 2 :]
    if remainder:
        segments.append((remainder, False))
    return segments


def _write_runs(paragraph, text: str, *, size_pt: int, italic: bool = False) -> None:
    """Lay `text` into a paragraph, giving every placeholder a run of its own.

    A placeholder has to be its own run so the renderer's run-level substitution can swap it
    whole and keep the surrounding formatting. Splitting here is what makes that true.
    """
    for piece, is_placeholder in _split_placeholders(text):
        run = paragraph.add_run(piece)
        house_style._set_font(run, size_pt=size_pt, bold=is_placeholder)
        run.italic = italic


def _build_line_items(document, table: LineItemTable) -> None:
    """The description/quantity/unit-price/amount table, as the reference invoices set it.

    Not a bordered grid. The reference rules the table in three places only - above the column
    headings, below them, and above the total - and leaves the body rows open. The headings run
    over two lines with the unit in brackets underneath, numeric columns are right-aligned and
    the description column is left-aligned, and the total row carries its label on the left with
    the figure hard against the right margin.

    One line item. Every reference invoice states a single commodity line, which is also all a
    transaction on this platform can carry: it has one commodity and one quantity. A multi-line
    invoice would need a data model that holds several, and there is no evidence in the
    references or in the platform that AGFZE raises one.
    """
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.shared import Inches

    grid = document.add_table(rows=3, cols=len(table.columns))
    grid.alignment = WD_TABLE_ALIGNMENT.LEFT
    grid.autofit = False

    # The description column takes the slack; the reference gives the three numeric columns
    # roughly equal, narrower widths and lets the description run wide.
    numeric_width = 1.15
    description_width = USABLE_WIDTH_IN - numeric_width * (len(table.columns) - 1)
    widths = [description_width] + [numeric_width] * (len(table.columns) - 1)

    heading_row, unit_row, body_row = grid.rows

    for index, column in enumerate(table.columns):
        alignment = WD_ALIGN_PARAGRAPH.RIGHT if column.numeric else WD_ALIGN_PARAGRAPH.LEFT
        for cell in (
            heading_row.cells[index],
            unit_row.cells[index],
            body_row.cells[index],
        ):
            cell.width = Inches(widths[index])

        head = heading_row.cells[index].paragraphs[0]
        head.alignment = alignment
        head.paragraph_format.space_after = Inches(0)
        house_style._set_font(
            head.add_run(column.heading), size_pt=house_style.SIZE_BODY_PT, bold=True
        )

        unit = unit_row.cells[index].paragraphs[0]
        unit.alignment = alignment
        unit.paragraph_format.space_after = Inches(0)
        _write_runs(unit, column.unit, size_pt=house_style.SIZE_BODY_PT)

        body = body_row.cells[index].paragraphs[0]
        body.alignment = alignment
        body.paragraph_format.space_after = Inches(0)
        _write_runs(
            body, PLACEHOLDER.format(name=column.field), size_pt=house_style.SIZE_BODY_PT
        )

    # The reference's three rules: above the headings, below them, and above the total.
    house_style.add_row_rule(heading_row, above=True)
    house_style.add_row_rule(unit_row, below=True)

    total = document.add_paragraph()
    house_style.add_rule(total, above=True)
    total.paragraph_format.tab_stops.add_tab_stop(Inches(USABLE_WIDTH_IN), WD_TAB_ALIGNMENT.RIGHT)
    _write_runs(total, table.total_label, size_pt=house_style.SIZE_BODY_PT)
    total.add_run("\t")
    _write_runs(
        total, PLACEHOLDER.format(name=table.total_field), size_pt=house_style.SIZE_BODY_PT
    )


def build_document(template: DocumentTemplate):
    """Write one template as a real, well-structured DOCX, in the AGFZE house style.

    Deterministic and reproducible: run it twice and you get the same document. Nothing here
    asks a model for anything - the template is the platform's own paperwork.

    The layout is the one measured from the approved reference documents in `house_style`:
    letterhead mark in the section header so it repeats on every page, registered-entity block in
    the section footer likewise, US Letter with the reference's asymmetric margins, and a body
    that runs title, reference line, counterparty block, metadata, line items, numbered clauses
    and signature block, in that order.

    One structural invariant matters and is load-bearing: the renderer pairs a clause's heading
    paragraph with *the paragraph immediately beneath it*. Every clause below is written as
    exactly that pair, and nothing is ever emitted between them.
    """
    from docx import Document as DocxDocument
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
    from docx.shared import Inches

    document = DocxDocument()

    house_style.apply_page_setup(document)
    house_style.apply_base_font(document)
    house_style.build_header(document)
    house_style.build_footer(document)

    body = document.add_paragraph()
    body.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = body.add_run(template.title)
    house_style._set_font(title, size_pt=house_style.SIZE_TITLE_PT, bold=True)
    title.underline = True

    # Not from the references, and deliberately so: the references are issued documents and this
    # is not. The draft marker is the platform's own, and it stays.
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    marker = subtitle.add_run(template.subtitle)
    house_style._set_font(marker, size_pt=house_style.SIZE_FOOTER_PT)
    marker.italic = True

    if template.reference_field and template.date_field:
        line = document.add_paragraph()
        line.paragraph_format.tab_stops.add_tab_stop(Inches(USABLE_WIDTH_IN), WD_TAB_ALIGNMENT.RIGHT)
        _write_runs(
            line,
            PLACEHOLDER.format(name=template.reference_field),
            size_pt=house_style.SIZE_REFERENCE_PT,
        )
        line.add_run("\t")
        _write_runs(
            line,
            "DATE: " + PLACEHOLDER.format(name=template.date_field),
            size_pt=house_style.SIZE_REFERENCE_PT,
        )

    if template.party_fields:
        document.add_paragraph()
        for name in template.party_fields:
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_after = Inches(0)
            _write_runs(
                paragraph, PLACEHOLDER.format(name=name), size_pt=house_style.SIZE_BODY_PT
            )

    # Everything a structural block has not already placed, in the reference's `LABEL : value`
    # pattern. Borderless: no reference document puts its fields in a ruled grid.
    placed = template.laid_out_field_names
    remaining = [item for item in template.fields if item.name not in placed]
    if remaining:
        document.add_paragraph()
        for item in remaining:
            paragraph = house_style.add_metadata_row(document, item.label.upper(), "")
            _write_runs(
                paragraph, PLACEHOLDER.format(name=item.name), size_pt=house_style.SIZE_BODY_PT
            )

    if template.line_items is not None:
        document.add_paragraph()
        _build_line_items(document, template.line_items)

    document.add_paragraph()

    for clause in template.clauses:
        clause_heading = document.add_paragraph()
        clause_heading.paragraph_format.space_after = Inches(0)
        # The marker identifies the clause to the renderer and is stripped from every produced
        # document, kept or revised. It never survives into a draft anybody reads.
        clause_heading.add_run(CLAUSE_MARKER.format(key=clause.key))
        # Upper case, because every reference document sets its clause headings that way. The
        # clause *number* is left exactly as the declaration writes it and is deliberately not
        # computed from position here: clauses are removed at render time, not build time, so a
        # position counted now would be wrong on any document that drops one. The declaration
        # numbers mutually exclusive alternates identically - the two price clauses are both
        # "4.", the two payment clauses both "6." - precisely so that whichever survives carries
        # the right number. The reference workbook solves the same problem by keeping a separate
        # sheet per pricing variant.
        house_style._set_font(
            clause_heading.add_run(clause.heading.upper()),
            size_pt=house_style.SIZE_BODY_PT,
            bold=True,
        )

        clause_body = document.add_paragraph()
        clause_body.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        _write_runs(clause_body, clause.body, size_pt=house_style.SIZE_BODY_PT)

    if template.signature_columns:
        house_style.add_signature_block(
            document,
            [
                (column.role, "FOR " + PLACEHOLDER.format(name=column.entity_field))
                for column in template.signature_columns
            ],
        )
        # Same rewrite as the reference line: the block is written with plain runs, and the
        # entity runs have to become placeholder runs for the renderer to fill them.
        signature_table = document.tables[-1]
        for cell, column in zip(signature_table.rows[0].cells, template.signature_columns):
            entity_paragraph = cell.paragraphs[1]
            for run in list(entity_paragraph.runs):
                run.text = ""
            _write_runs(
                entity_paragraph,
                "FOR " + PLACEHOLDER.format(name=column.entity_field),
                size_pt=house_style.SIZE_BODY_PT,
            )

    document.add_paragraph()
    footer = document.add_paragraph()
    _write_runs(footer, template.footer_note, size_pt=house_style.SIZE_FOOTER_PT, italic=True)

    return document


def ensure_template_files(*, overwrite: bool = False) -> BuildReport:
    """Materialise every declared template on disk, so the renderer always has a file to open.

    Called once at application start and by the build entry point. Idempotent: an existing
    template is left exactly as it is unless `overwrite` is asked for, so an administrator who
    has adjusted the shipped wording does not have it silently replaced on the next deploy.
    """
    ASSET_ROOT.mkdir(parents=True, exist_ok=True)
    report = BuildReport()
    for template in TEMPLATES_BY_TYPE.values():
        destination = template_path(template)
        if destination.exists() and not overwrite:
            report.skipped.append(destination.name)
            continue
        build_document(template).save(str(destination))
        report.written.append(destination.name)
    return report


if __name__ == "__main__":  # pragma: no cover - a build entry point, not application code
    result = ensure_template_files(overwrite=True)
    for name in result.written:
        print(f"wrote {name}")
