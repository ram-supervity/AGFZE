"""Deterministic reading of a document before any model is involved.

A digital PDF carries its own text layer, and reading it with PyMuPDF is both cheaper and more
faithful than asking a vision model to read a picture of it. Only when that layer is absent or
too thin to be real - a scan, a photograph - does the page get rasterised for the multimodal
path. Word, Excel and CSV are always deterministic; there is nothing for a vision model to add.
"""

from __future__ import annotations

import csv
import io
import unicodedata
from dataclasses import dataclass, field
from enum import Enum

# PyMuPDF's historical import name `fitz` is deprecated from 1.26 onwards.
import pandas as pd
import pymupdf
from docx import Document as DocxDocument

from app.core.config import settings
from app.core.logging import get_logger
from app.services.file_intake import DOCX, IMAGE, PDF, SPREADSHEET

logger = get_logger(__name__)

class ExtractionRoute(str, Enum):
    TEXT_LAYER = "text_layer"
    MULTIMODAL = "multimodal"
    OFFICE = "office"
    TABULAR = "tabular"


@dataclass
class PageContent:
    page_number: int
    text: str
    # Paragraph ordinal -> its text, so a field can point back at the exact block it was read from.
    blocks: list[dict[str, object]] = field(default_factory=list)
    image: bytes | None = None
    image_mime: str = "image/png"


@dataclass
class DocumentContent:
    route: ExtractionRoute
    pages: list[PageContent]

    @property
    def page_count(self) -> int:
        return len(self.pages)

    @property
    def full_text(self) -> str:
        return "\n\n".join(page.text for page in self.pages if page.text.strip())


def normalise_text(value: str) -> str:
    """Fold full-width digits, CJK punctuation and compatibility forms to a comparable form.

    A China-territory certificate quotes its contract number in full-width characters; the same
    reference on the invoice is plain ASCII. NFKC makes the two comparable without discarding
    the non-Latin body text around them.
    """
    if not value:
        return ""
    folded = unicodedata.normalize("NFKC", value)
    return "".join(" " if character in "　\xa0" else character for character in folded)


def _pdf_page_text(page: pymupdf.Page) -> tuple[str, list[dict[str, object]]]:
    blocks: list[dict[str, object]] = []
    raw = page.get_text("blocks") or []
    lines: list[str] = []
    for ordinal, block in enumerate(sorted(raw, key=lambda item: (round(item[1], 1), item[0]))):
        text = normalise_text(str(block[4] or "")).strip()
        if not text:
            continue
        lines.append(text)
        blocks.append(
            {
                "paragraph": ordinal,
                "bbox": [round(float(value), 2) for value in block[:4]],
                "text": text[:400],
            }
        )
    return "\n".join(lines), blocks


def _rasterise(page: pymupdf.Page, dpi: int) -> bytes:
    pixmap = page.get_pixmap(dpi=dpi)
    return pixmap.tobytes("png")


def read_pdf(data: bytes) -> DocumentContent:
    """Read a PDF, converting each page into an image for multimodal AI processing."""
    pages: list[PageContent] = []
    with pymupdf.open(stream=data, filetype="pdf") as document:
        limit = min(document.page_count, settings.EXTRACTION_MAX_PAGES)
        for index in range(limit):
            page = document.load_page(index)
            text, blocks = _pdf_page_text(page)
            image = _rasterise(page, settings.PAGE_RASTER_DPI)
            pages.append(
                PageContent(
                    page_number=index + 1,
                    text=text,
                    blocks=blocks,
                    image=image,
                    image_mime="image/png",
                )
            )

    return DocumentContent(route=ExtractionRoute.MULTIMODAL, pages=pages)


def render_pdf_pages(data: bytes, page_numbers: list[int] | None = None) -> list[bytes]:
    """Rasterise pages for the review viewer when the text-layer path skipped rendering."""
    rendered: list[bytes] = []
    with pymupdf.open(stream=data, filetype="pdf") as document:
        limit = min(document.page_count, settings.EXTRACTION_MAX_PAGES)
        wanted = page_numbers or list(range(1, limit + 1))
        for page_number in wanted:
            if 1 <= page_number <= limit:
                page = document.load_page(page_number - 1)
                rendered.append(_rasterise(page, settings.PAGE_RASTER_DPI))
    return rendered


def render_document_preview_pages(filename: str, content: DocumentContent, dpi: int = 150) -> list[bytes]:
    """Generate rendered PNG page previews for any document content."""
    rendered_images: list[bytes] = []
    pages = content.pages or [PageContent(page_number=1, text="")]
    for page in pages:
        if page.image is not None:
            rendered_images.append(page.image)
            continue

        doc = pymupdf.open()
        p = doc.new_page(width=595, height=842)
        p.insert_text((40, 40), filename, fontsize=11, fontname="helv", color=(0.2, 0.2, 0.2))
        page_label = f"Page {page.page_number}"
        text_len = pymupdf.get_text_length(page_label, fontname="helv", fontsize=9)
        p.insert_text((555 - text_len, 40), page_label, fontsize=9, fontname="helv", color=(0.4, 0.4, 0.4))
        p.draw_line(pymupdf.Point(40, 48), pymupdf.Point(555, 48), color=(0.8, 0.8, 0.8), width=0.6)

        text = page.text if hasattr(page, "text") and page.text else "(empty document)"
        rect = pymupdf.Rect(40, 60, 555, 800)
        p.insert_textbox(rect, text, fontsize=9.5, fontname="helv", color=(0.1, 0.1, 0.1))

        pix = p.get_pixmap(dpi=dpi)
        rendered_images.append(pix.tobytes("png"))

    return rendered_images


def read_image(data: bytes, mime: str) -> DocumentContent:
    """A photograph or a scan handed in directly: one page, read by the multimodal path."""
    return DocumentContent(
        route=ExtractionRoute.MULTIMODAL,
        pages=[PageContent(page_number=1, text="", image=data, image_mime=mime)],
    )


def read_docx(data: bytes) -> DocumentContent:
    document = DocxDocument(io.BytesIO(data))
    lines: list[str] = []
    blocks: list[dict[str, object]] = []

    for ordinal, paragraph in enumerate(document.paragraphs):
        text = normalise_text(paragraph.text).strip()
        if not text:
            continue
        lines.append(text)
        blocks.append({"paragraph": ordinal, "text": text[:400]})

    for table_index, table in enumerate(document.tables):
        for row_index, row in enumerate(table.rows):
            cells = [normalise_text(cell.text).strip() for cell in row.cells]
            if not any(cells):
                continue
            text = " | ".join(cells)
            lines.append(text)
            blocks.append(
                {
                    "paragraph": len(blocks),
                    "table": table_index,
                    "row": row_index,
                    "text": text[:400],
                }
            )

    return DocumentContent(
        route=ExtractionRoute.OFFICE,
        pages=[PageContent(page_number=1, text="\n".join(lines), blocks=blocks)],
    )


def _detect_header_row(frame: pd.DataFrame) -> int:
    """Tracker exports carry a title banner above the real header; find the header row.

    The header is taken to be the first row where most cells are non-empty, distinct strings.
    """
    for index in range(min(len(frame), 15)):
        values = [str(value).strip() for value in frame.iloc[index].tolist()]
        populated = [value for value in values if value and value.lower() != "nan"]
        if len(populated) >= max(2, len(values) // 2) and len(set(populated)) == len(populated):
            return index
    return 0


def _frame_to_text(frame: pd.DataFrame, sheet_name: str | None) -> tuple[str, list[dict]]:
    header_row = _detect_header_row(frame)
    header = [normalise_text(str(value)).strip() for value in frame.iloc[header_row].tolist()]
    body = frame.iloc[header_row + 1 :]

    lines: list[str] = []
    blocks: list[dict[str, object]] = []
    if sheet_name:
        lines.append(f"# sheet: {sheet_name}")
    lines.append(" | ".join(header))

    for offset, (_, row) in enumerate(body.iterrows()):
        cells = [normalise_text(str(value)).strip() for value in row.tolist()]
        if not any(cell and cell.lower() != "nan" for cell in cells):
            continue
        rendered = " | ".join("" if cell.lower() == "nan" else cell for cell in cells)
        lines.append(rendered)
        blocks.append(
            {
                "paragraph": len(blocks),
                "sheet": sheet_name,
                "row": header_row + 1 + offset,
                "text": rendered[:400],
            }
        )

    return "\n".join(lines), blocks


def _read_csv_frame(data: bytes) -> pd.DataFrame:
    """Read a CSV whose rows are not all the same width.

    A tracker export routinely opens with a one-cell title banner above the real header, which
    pandas' own parser rejects as malformed. Reading through the stdlib csv module and padding
    every row to the widest one keeps that file readable instead of failing on it.
    """
    text = data.decode("utf-8-sig", errors="replace")
    try:
        delimiter = csv.Sniffer().sniff(text[:4096], delimiters=",;\t|").delimiter
    except csv.Error:
        delimiter = ","
    rows = list(csv.reader(io.StringIO(text), delimiter=delimiter))
    width = max((len(row) for row in rows), default=0)
    if width == 0:
        return pd.DataFrame()
    padded = [row + [""] * (width - len(row)) for row in rows]
    return pd.DataFrame(padded, dtype=str).fillna("")


def read_spreadsheet(data: bytes, content_type: str) -> DocumentContent:
    """Parse a tracker export tolerantly: no assumed header position, no assumed sheet name."""
    pages: list[PageContent] = []

    if content_type == "text/csv":
        page_text, blocks = _frame_to_text(_read_csv_frame(data), None)
        pages.append(PageContent(page_number=1, text=page_text, blocks=blocks))
    else:
        engine = "openpyxl" if "spreadsheetml" in content_type else None
        sheets = pd.read_excel(
            io.BytesIO(data), header=None, dtype=str, sheet_name=None, engine=engine
        )
        for index, (sheet_name, frame) in enumerate(sheets.items(), start=1):
            page_text, blocks = _frame_to_text(frame.fillna(""), str(sheet_name))
            pages.append(PageContent(page_number=index, text=page_text, blocks=blocks))

    if not pages:
        pages.append(PageContent(page_number=1, text=""))
    return DocumentContent(route=ExtractionRoute.TABULAR, pages=pages)


def read_document(data: bytes, *, family: str, content_type: str) -> DocumentContent:
    """Route a file to the reader its real bytes call for.

    `family` comes from libmagic via `file_intake.detect_type`, never from the filename.
    """
    if family == PDF:
        return read_pdf(data)
    if family == IMAGE:
        return read_image(data, content_type)
    if family == DOCX:
        return read_docx(data)
    if family == SPREADSHEET:
        return read_spreadsheet(data, content_type)
    raise ValueError(f"No reader is registered for file family: {family}")
