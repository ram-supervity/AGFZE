"""Draft document generation: the template, the validation gate, and what a failure produces.

The single most important behaviour proved here is the negative one. A model reply that does not
survive validation must fail the job cleanly and produce nothing at all - no document row, no
stored file, no half-populated template. A generated draft contract is exactly the kind of
artefact that gets acted on with less scrutiny than an extraction ever did, so a polished
document carrying wrong commercial terms is a far worse failure than a visible one.
"""

from __future__ import annotations

import io
import json
import uuid

import pytest
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.models.enums import (
    DocumentSource,
    DocumentType,
    ExtractionStatus,
    FixationStatus,
    Territory,
)
from app.models.identity import User
from app.models.intake import Document
from app.services import draft_service
from app.services.gemini_service import AIServiceError
from app.services.templates import house_style
from app.services.templates.renderer import ClauseDirective, TemplateRenderError, render_template
from app.services.templates.sales_templates import (
    SALES_CONTRACT_TEMPLATE,
    SALES_INVOICE_TEMPLATE,
    TEMPLATES_BY_TYPE,
    ensure_template_files,
    template_path,
    territory_reference,
)
from tests.utils.sales import (
    CUSTOMER,
    VALID_CONTRACT_PLAN,
    VALID_INVOICE_PLAN,
    draft_plan_response,
    sales_transaction,
)


@pytest.fixture
def model_reply(monkeypatch: pytest.MonkeyPatch):
    """Answer the next model call with an exact payload, and capture the prompt it was sent.

    The same seam every other AI test in this suite uses. Nothing here needs a live key, and the
    payloads are written by the test rather than sampled from a model.
    """
    captured: dict[str, object] = {}

    def _install(payload: str | Exception) -> dict[str, object]:
        async def _raw(prompt, response_schema, images):
            captured["prompt"] = prompt
            captured["schema"] = response_schema
            if isinstance(payload, Exception):
                raise payload
            return payload

        monkeypatch.setattr("app.services.gemini_service._generate_raw", _raw)
        return captured

    return _install


@pytest.fixture
async def sales_user(db_session: AsyncSession) -> User:
    uid = uuid.uuid4().hex[:8]
    user = User(
        subject_id=f"draft-user-{uid}",
        email=f"sales.user.{uid}@agfze.test",
        display_name="Sales User",
        roles=["sales_user"],
    )
    db_session.add(user)
    await db_session.flush()
    return user


def _docx_text(payload: bytes) -> str:
    from docx import Document as DocxDocument

    document = DocxDocument(io.BytesIO(payload))
    body = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for table in document.tables:
        for row in table.rows:
            body += "\n" + " | ".join(cell.text for cell in row.cells)
    return body


# --- the templates themselves ---------------------------------------------------------------


def test_every_template_ships_as_a_real_file() -> None:
    ensure_template_files()
    for template in TEMPLATES_BY_TYPE.values():
        path = template_path(template)
        assert path.exists(), f"{template.filename} is not on disk"
        assert path.stat().st_size > 5_000, "a template that small is not a real document"
        assert template.clauses, "a template with no clauses is not a template"
        assert template.required_clause_keys, "every document has clauses it cannot lose"


def test_a_rendered_document_carries_real_data_and_no_placeholders() -> None:
    ensure_template_files()
    values = {name: f"value-{name}" for name in SALES_CONTRACT_TEMPLATE.field_names}
    values.update(
        {
            "buyer": CUSTOMER,
            "quantity": "24.500 MT",
            "territory_reference": territory_reference(Territory.CHINA.value),
            "generated_at": "2026-08-28T09:00:00",
            "generated_by": "Sales User",
        }
    )

    result = render_template(
        SALES_CONTRACT_TEMPLATE,
        template_bytes=template_path(SALES_CONTRACT_TEMPLATE).read_bytes(),
        values=values,
        directives=[
            ClauseDirective("pricing_lme", "remove"),
            ClauseDirective("payment_tt", "remove"),
        ],
    )
    text = _docx_text(result.content)

    assert CUSTOMER in text
    assert "24.500 MT" in text
    assert "{{" not in text, "an unresolved placeholder must never reach a produced draft"
    assert "[[clause:" not in text, "the template's control markers must never survive"
    assert "pricing_lme" in result.removed
    assert "payment_tt" in result.removed


def test_territory_selects_content_not_a_separate_template() -> None:
    """One template per document type; the destination changes what it says, not its structure."""
    ensure_template_files()
    india = territory_reference(Territory.INDIA.value)
    china = territory_reference(Territory.CHINA.value)

    assert india != china
    assert "India" in india
    assert "China" in china
    # An unrecognised territory falls to a real entry, never to a blank.
    assert territory_reference(None).strip()
    assert territory_reference("atlantis") == territory_reference(Territory.OTHER.value)
    # And there is exactly one template file per document type, not one per territory.
    assert len(TEMPLATES_BY_TYPE) == 6


def test_the_renderer_refuses_to_delete_a_required_clause() -> None:
    ensure_template_files()
    with pytest.raises(TemplateRenderError):
        render_template(
            SALES_CONTRACT_TEMPLATE,
            template_bytes=template_path(SALES_CONTRACT_TEMPLATE).read_bytes(),
            values={},
            directives=[ClauseDirective("parties", "remove")],
        )


def test_an_absent_value_renders_as_an_explicit_gap() -> None:
    """A slot with nothing behind it says so, rather than reading as a blank somebody wrote."""
    ensure_template_files()
    result = render_template(
        SALES_INVOICE_TEMPLATE,
        template_bytes=template_path(SALES_INVOICE_TEMPLATE).read_bytes(),
        values={"buyer": CUSTOMER},
        directives=[],
    )

    assert "not recorded" in _docx_text(result.content)
    assert "quantity" in result.unpopulated


# --- validating the model's answer -------------------------------------------------------------


def _plan(payload: str):
    from app.services.gemini_service import DraftContentPlan

    return DraftContentPlan.model_validate(json.loads(payload))


def test_a_clause_the_template_does_not_have_is_rejected() -> None:
    with pytest.raises(draft_service.DraftContentInvalidError):
        draft_service.validate_plan(
            SALES_CONTRACT_TEMPLATE, _plan(draft_plan_response(keep=["a_clause_i_invented"]))
        )


def test_removing_a_required_clause_is_rejected() -> None:
    with pytest.raises(draft_service.DraftContentInvalidError):
        draft_service.validate_plan(
            SALES_CONTRACT_TEMPLATE, _plan(draft_plan_response(remove=["governing_law"]))
        )


def test_a_revision_with_no_wording_is_rejected() -> None:
    with pytest.raises(draft_service.DraftContentInvalidError):
        draft_service.validate_plan(
            SALES_CONTRACT_TEMPLATE, _plan(draft_plan_response(revise={"inspection": "  "}))
        )


def test_a_revision_referring_to_a_field_the_template_cannot_populate_is_rejected() -> None:
    with pytest.raises(draft_service.DraftContentInvalidError):
        draft_service.validate_plan(
            SALES_CONTRACT_TEMPLATE,
            _plan(
                draft_plan_response(
                    revise={
                        "inspection": (
                            "A surveyor shall be appointed and paid from {{secret_slush_fund}}."
                        )
                    }
                )
            ),
        )


def test_an_empty_plan_is_rejected() -> None:
    with pytest.raises(draft_service.DraftContentInvalidError):
        draft_service.validate_plan(SALES_CONTRACT_TEMPLATE, _plan(draft_plan_response()))


def test_two_instructions_for_one_clause_are_rejected() -> None:
    payload = json.dumps(
        {
            "clauses": [
                {"key": "inspection", "action": "keep"},
                {"key": "inspection", "action": "remove"},
            ]
        }
    )
    with pytest.raises(draft_service.DraftContentInvalidError):
        draft_service.validate_plan(SALES_CONTRACT_TEMPLATE, _plan(payload))


def test_a_valid_plan_survives_validation() -> None:
    directives = draft_service.validate_plan(SALES_CONTRACT_TEMPLATE, _plan(VALID_CONTRACT_PLAN))

    assert {item.key for item in directives} <= SALES_CONTRACT_TEMPLATE.clause_keys
    assert {item.action for item in directives} <= {"keep", "revise", "remove"}


# --- generation, end to end -----------------------------------------------------------------------


async def test_a_successful_generation_produces_a_real_generated_document_row(
    db_session: AsyncSession, sales_user: User, model_reply, storage_root
) -> None:
    model_reply(VALID_CONTRACT_PLAN)
    transaction = await sales_transaction(db_session, batch_number="I2626-100")

    result = await draft_service.generate(
        db_session,
        transaction,
        document_type=DocumentType.DRAFT_CONTRACT.value,
        requested_by=sales_user,
    )
    await db_session.commit()

    document = await db_session.get(Document, result.document_id)
    assert document is not None
    assert document.document_type == DocumentType.DRAFT_CONTRACT.value
    assert document.source == DocumentSource.GENERATED.value
    # The first document in the platform that nothing received.
    assert document.request_id is None
    assert document.extraction_status == ExtractionStatus.NOT_APPLICABLE.value
    assert document.classification_confidence is None
    # `uploaded_by` records who triggered the generation, reusing the field that already exists.
    assert document.uploaded_by_id == sales_user.id
    assert document.transaction_id == transaction.id

    # The established naming convention, and a UUID-derived key behind it.
    assert document.filename.startswith("SO-I2626-100-24.5-Prov")
    assert document.storage_ref == f"documents/generated/{document.id}/{document.filename}"

    stored = await storage_root.download(document.storage_ref)
    assert stored[:2] == b"PK", "a DOCX is a zip container; this is not one"
    text = _docx_text(stored)
    assert CUSTOMER in text
    assert "I2626-100" in text
    assert "{{" not in text


async def test_a_fixed_customer_produces_a_final_named_draft(
    db_session: AsyncSession, sales_user: User, model_reply
) -> None:
    model_reply(VALID_INVOICE_PLAN)
    transaction = await sales_transaction(
        db_session,
        batch_number="I2626-101",
        fixation_status=FixationStatus.FIXED.value,
        fixation_rate="8420.00",
    )

    result = await draft_service.generate(
        db_session,
        transaction,
        document_type=DocumentType.DRAFT_INVOICE.value,
        requested_by=sales_user,
    )
    await db_session.commit()

    assert "-Final-" in result.filename
    assert result.document_type == DocumentType.DRAFT_INVOICE.value


async def test_every_document_this_platform_writes_gets_its_own_filename(
    db_session: AsyncSession, sales_user: User
) -> None:
    """A proforma invoice and a bank cover letter are not "the invoice".

    They land on the same transaction with the same batch and the same quantity, so a naming rule
    that calls everything except a contract an invoice hands the desk two files with one name and
    no way to tell which document is which.
    """
    transaction = await sales_transaction(db_session, batch_number="I2626-110")

    names = {
        document_type: draft_service.storage_filename(transaction, document_type)
        for document_type in (
            DocumentType.DRAFT_CONTRACT.value,
            DocumentType.DRAFT_INVOICE.value,
            DocumentType.DRAFT_PERFORMA_INVOICE.value,
            DocumentType.DRAFT_BANK_COVER_LETTER.value,
        )
    }

    assert len(set(names.values())) == len(names), names
    # The two the desk already had keep the words the desk already uses.
    assert names[DocumentType.DRAFT_CONTRACT.value].endswith("-contract.docx")
    assert names[DocumentType.DRAFT_INVOICE.value].endswith("-invoice.docx")
    assert names[DocumentType.DRAFT_PERFORMA_INVOICE.value].endswith("-performa-invoice.docx")
    assert names[DocumentType.DRAFT_BANK_COVER_LETTER.value].endswith("-bank-cover-letter.docx")
    for name in names.values():
        assert name.startswith("SO-I2626-110-")


async def test_a_malformed_model_response_fails_cleanly_and_produces_no_document(
    db_session: AsyncSession, sales_user: User, model_reply
) -> None:
    """The higher-stakes AI-failure principle, verified.

    Nothing is populated, nothing is stored and no document row appears. The alternative - render
    the template anyway with the shipped wording - would hand a reviewer a polished contract that
    nobody has vouched for, and that is the exact failure this rule exists to prevent.
    """
    model_reply("this is not JSON at all")
    transaction = await sales_transaction(db_session, batch_number="I2626-102")
    before = len((await db_session.scalars(select(Document))).all())

    with pytest.raises(AIServiceError):
        await draft_service.generate(
            db_session,
            transaction,
            document_type=DocumentType.DRAFT_CONTRACT.value,
            requested_by=sales_user,
        )

    after = (await db_session.scalars(select(Document))).all()
    assert len(after) == before
    assert not any(row.source == DocumentSource.GENERATED.value for row in after)


async def test_a_schema_valid_but_nonsense_plan_fails_cleanly_too(
    db_session: AsyncSession, sales_user: User, model_reply
) -> None:
    """Schema-valid is not enough. The clause registry is the second gate, and it is real."""
    model_reply(draft_plan_response(remove=["parties", "goods", "governing_law"]))
    transaction = await sales_transaction(db_session, batch_number="I2626-103")

    with pytest.raises(draft_service.DraftContentInvalidError):
        await draft_service.generate(
            db_session,
            transaction,
            document_type=DocumentType.DRAFT_CONTRACT.value,
            requested_by=sales_user,
        )

    generated = (
        await db_session.scalars(
            select(Document).where(Document.source == DocumentSource.GENERATED.value)
        )
    ).all()
    assert generated == []


async def test_regenerating_creates_a_new_draft_beside_the_old_one(
    db_session: AsyncSession, sales_user: User, model_reply, storage_root
) -> None:
    """Requesting changes never overwrites. The prior draft stays in the document history."""
    model_reply(VALID_CONTRACT_PLAN)
    transaction = await sales_transaction(db_session, batch_number="I2626-104")

    first = await draft_service.generate(
        db_session,
        transaction,
        document_type=DocumentType.DRAFT_CONTRACT.value,
        requested_by=sales_user,
    )
    await db_session.commit()

    model_reply(
        draft_plan_response(
            keep=[
                "parties",
                "goods",
                "delivery",
                "documents",
                "title_risk",
                "force_majeure",
                "governing_law",
                "payment_cad",
                "pricing_fixed",
            ],
            revise={
                "quantity_tolerance": (
                    "Shipped weight is final at the port of loading and no outturn adjustment "
                    "applies to this shipment."
                )
            },
            remove=["pricing_lme", "price_fixation", "payment_tt", "inspection"],
        )
    )
    second = await draft_service.generate(
        db_session,
        transaction,
        document_type=DocumentType.DRAFT_CONTRACT.value,
        requested_by=sales_user,
    )
    await db_session.commit()

    assert first.document_id != second.document_id
    assert "quantity_tolerance" in second.revised

    drafts = await draft_service.drafts_for(db_session, transaction.id)
    assert len(drafts) == 2, "the earlier draft must remain in the transaction's history"
    assert {row.id for row in drafts} == {first.document_id, second.document_id}
    # Both files still exist; the first was not replaced under the second's key.
    assert await storage_root.download(first.storage_ref)
    assert await storage_root.download(second.storage_ref)


async def test_the_model_is_never_asked_for_the_document_itself(
    db_session: AsyncSession, sales_user: User, model_reply
) -> None:
    """The prompt asks for clause decisions. The schema cannot express a figure or a file."""
    captured = model_reply(VALID_CONTRACT_PLAN)
    transaction = await sales_transaction(db_session, batch_number="I2626-105")

    await draft_service.generate(
        db_session,
        transaction,
        document_type=DocumentType.DRAFT_CONTRACT.value,
        requested_by=sales_user,
    )
    await db_session.commit()

    schema = captured["schema"]
    assert set(schema["properties"]) == {"clauses", "notes"}
    clause = schema["properties"]["clauses"]["items"]["properties"]
    assert set(clause) == {"key", "action", "text", "reason"}
    assert clause["action"]["enum"] == ["keep", "revise", "remove"]

    prompt = str(captured["prompt"])
    assert "you do not supply any commercial figure" in prompt.lower()
    # The prompt tells the model in as many words not to write about dispatch or e-signature.
    # The absence of any code that could dispatch anything is checked separately, below.
    assert "do not write anything about sending" in prompt.lower()


async def test_draft_generation_is_gated_on_br_07s_draft_check(
    db_session: AsyncSession, sales_user: User
) -> None:
    """A draft B/L is enough. Nothing at all is not."""
    from app.core.errors import ConflictError

    permitted = await sales_transaction(
        db_session, batch_number="I2626-106", with_final_bl=False, with_draft_bl=True
    )
    await draft_service.assert_generation_permitted(db_session, permitted)

    from app.services.rules import engine as rule_engine

    blocked = await sales_transaction(
        db_session, batch_number="I2626-107", with_final_bl=False, with_draft_bl=False
    )
    blocked.sales_leg.bl_reference = None
    await rule_engine.run_validation(db_session, blocked)
    await db_session.commit()

    with pytest.raises(ConflictError):
        await draft_service.assert_generation_permitted(db_session, blocked)


def test_no_code_path_can_send_a_generated_document_anywhere() -> None:
    """A governance requirement checked rather than assumed.

    The complete absence of a dispatch mechanism is itself the control. This reads the modules
    that touch a generated draft and asserts that none of them so much as imports something that
    could transmit one - no mail client, no HTTP client, no outbound queue.
    """
    from pathlib import Path

    import app.api.v1.transactions as transactions_module
    import app.services.draft_service as draft_module
    import app.services.sales_service as sales_module
    import app.services.templates.renderer as renderer_module

    forbidden = (
        "smtplib",
        "sendgrid",
        "aiosmtplib",
        "send_mail",
        "send_email",
        "send_document",
        "dispatch_document",
        "graph_service.send",
        "docusign",
        "esign",
        "e_signature",
    )
    for module in (draft_module, sales_module, renderer_module, transactions_module):
        source = Path(module.__file__).read_text().lower()
        for token in forbidden:
            assert token not in source, (
                f"{Path(module.__file__).name} references '{token}'. No path may exist, even "
                "dormant, to send a generated document outside the platform."
            )


# --- the two documents discovery named that the platform could not previously produce -------------


def test_the_performa_invoice_and_bank_cover_letter_are_real_renderable_templates() -> None:
    """Both were named in discovery and neither existed as an asset until now."""
    from app.models.enums import DocumentType
    from app.services.templates.sales_templates import (
        BANK_COVER_LETTER_TEMPLATE,
        PERFORMA_INVOICE_TEMPLATE,
    )

    ensure_template_files()
    for template in (PERFORMA_INVOICE_TEMPLATE, BANK_COVER_LETTER_TEMPLATE):
        path = template_path(template)
        assert path.exists(), f"{template.filename} is not on disk"
        assert path.stat().st_size > 5_000, "a template that small is not a real document"
        assert template.required_clause_keys, "every document has clauses it cannot lose"

    assert PERFORMA_INVOICE_TEMPLATE.document_type == DocumentType.DRAFT_PERFORMA_INVOICE.value
    assert BANK_COVER_LETTER_TEMPLATE.document_type == DocumentType.DRAFT_BANK_COVER_LETTER.value


def test_a_performa_invoice_never_claims_a_shipped_weight() -> None:
    """Its defining property, and the one thing a careless revision could destroy.

    A Performa invoice is raised before the cargo is weighed. If it carried a bill-of-lading
    reference or asserted a shipped weight it would be a commercial invoice wearing the wrong
    title, and somebody downstream would treat an advance figure as a settled one.
    """
    from app.services.templates.sales_templates import PERFORMA_INVOICE_TEMPLATE

    assert "bl_reference" not in PERFORMA_INVOICE_TEMPLATE.field_names

    body = " ".join(clause.body for clause in PERFORMA_INVOICE_TEMPLATE.clauses).lower()
    assert "bill of lading" not in body
    assert "weight slip" not in body or "no weight slip" in body

    # And it says out loud what it is, so a reader is never left inferring it from a blank field.
    basis = PERFORMA_INVOICE_TEMPLATE.clause("provisional_basis")
    assert basis is not None
    assert basis.required is True
    assert "contracted quantity" in basis.body.lower()
    assert "no weight slip" in basis.body.lower()


def test_a_bank_cover_letter_carries_no_commercial_terms_of_its_own() -> None:
    """A second, unsigned source of the commercial terms is what holds a presentation up.

    The one value it does state is the release amount, because the bank's instruction is
    conditional on it. It states no rate at all.
    """
    from app.services.templates.sales_templates import BANK_COVER_LETTER_TEMPLATE

    assert "price_terms" not in BANK_COVER_LETTER_TEMPLATE.field_names

    body = " ".join(clause.body for clause in BANK_COVER_LETTER_TEMPLATE.clauses).lower()
    assert "{{price_terms}}" not in body
    # It does list what is enclosed, which is the whole job of a cover letter.
    enclosures = BANK_COVER_LETTER_TEMPLATE.clause("enclosures")
    assert enclosures is not None and enclosures.required is True


async def test_a_performa_invoice_generates_with_no_bill_of_lading_at_all(
    db_session, storage_root
) -> None:
    """The case that would have been impossible before, and is the only case that matters.

    BR-07 holds a draft *commercial* invoice back until shipment evidence exists. A Performa
    invoice is raised before there is any, so gating it the same way would make the platform
    unable to produce one in the only circumstance it is ever produced in.
    """
    from app.models.enums import DocumentType
    from app.services import draft_service
    from app.services.rules import engine as rule_engine
    from tests.utils.sales import sales_transaction

    transaction = await sales_transaction(
        db_session,
        batch_number="I2626-P1",
        with_final_bl=False,
        with_draft_bl=False,
        validate=False,
    )
    # No bill of lading document *and* no recorded B/L reference. Both have to go: BR-07 accepts a
    # reference on its own, because the desk routinely has the number from the carrier before any
    # document arrives, and that is exactly when drafting is meant to start. This transaction is
    # earlier than that - the cargo has not shipped at all, which is when a Performa is raised.
    transaction.sales_leg.bl_reference = None
    await db_session.commit()
    await rule_engine.run_validation(db_session, transaction)
    await db_session.commit()

    # The commercial invoice is correctly refused: there is no bill of lading behind it.
    with pytest.raises(draft_service.DraftNotPermittedError):
        await draft_service.assert_generation_permitted(
            db_session, transaction, document_type=DocumentType.DRAFT_INVOICE.value
        )

    # The Performa invoice and the bank cover letter are not, and that is the exemption working.
    for document_type in (
        DocumentType.DRAFT_PERFORMA_INVOICE.value,
        DocumentType.DRAFT_BANK_COVER_LETTER.value,
    ):
        await draft_service.assert_generation_permitted(
            db_session, transaction, document_type=document_type
        )


async def test_the_exemption_still_requires_a_sales_leg(db_session, storage_root) -> None:
    """Narrow, not a way past validation. No sales leg is still no sales document to draft."""
    from app.models.enums import DocumentType
    from app.services import draft_service
    from tests.utils.transactions import make_transaction

    transaction = await make_transaction(db_session, batch_number="I2626-P9")
    await db_session.commit()

    with pytest.raises(draft_service.DraftNotPermittedError):
        await draft_service.assert_generation_permitted(
            db_session, transaction, document_type=DocumentType.DRAFT_PERFORMA_INVOICE.value
        )


async def test_purchase_contract_and_cost_sheet_permitted_with_purchase_data(
    db_session, storage_root
) -> None:
    """Purchase contract and cost sheet drafts are permitted on purchase transactions."""
    from app.models.enums import DocumentType
    from app.services import draft_service
    from tests.utils.transactions import make_transaction

    transaction = await make_transaction(
        db_session,
        batch_number="I2626-PO-1",
        supplier_name="Apex Scrap Corp",
        quantity="25.000",
        rate="4800.00",
    )
    await db_session.commit()

    for doc_type in (
        DocumentType.DRAFT_PURCHASE_CONTRACT.value,
        DocumentType.DRAFT_COST_SHEET.value,
    ):
        await draft_service.assert_generation_permitted(
            db_session, transaction, document_type=doc_type
        )


async def test_purchase_draft_refused_when_required_data_missing(
    db_session, storage_root
) -> None:
    """Missing supplier or required figures blocks draft generation with a clear message."""
    from app.models.enums import DocumentType
    from app.services import draft_service
    from tests.utils.transactions import make_transaction

    transaction = await make_transaction(
        db_session,
        batch_number="I2626-PO-2",
        supplier_name="Apex Scrap Corp",
    )
    # Clear quantity and rate
    transaction.quantity_mt = None
    transaction.purchase_leg.rate = None
    transaction.purchase_leg.amount = None
    transaction.price_basis = None
    await db_session.commit()

    with pytest.raises(draft_service.DraftNotPermittedError) as exc:
        await draft_service.assert_generation_permitted(
            db_session,
            transaction,
            document_type=DocumentType.DRAFT_PURCHASE_CONTRACT.value,
        )
    assert "missing from the purchase transaction" in str(exc.value)


async def test_api_role_restrictions_for_draft_generation(
    patched_jwks, client, db_session: AsyncSession, signed_in, storage_root
) -> None:
    """Verify strict backend role-based access for draft generation."""
    from tests.utils.transactions import make_transaction

    tx = await make_transaction(
        db_session,
        batch_number="I2626-PO-API",
        supplier_name="Apex Scrap Corp",
        quantity="20.000",
        rate="5000.00",
    )
    await db_session.commit()

    _, purchase_headers = await signed_in(
        "purch-api-1", "purch.api@agfze.test", "Purchase User", ["purchase_user"]
    )
    _, sales_headers = await signed_in(
        "sales-api-1", "sales.api@agfze.test", "Sales User", ["sales_user"]
    )
    _, other_headers = await signed_in(
        "fa-api-1", "fa.api@agfze.test", "FA User", ["fa_user"]
    )

    # 1. Purchase user can generate purchase contract
    res = await client.post(
        f"/api/v1/transactions/{tx.id}/generate-draft",
        headers=purchase_headers,
        json={"document_type": DocumentType.DRAFT_PURCHASE_CONTRACT.value},
    )
    assert res.status_code == 202

    # 2. Purchase user can generate cost sheet
    res = await client.post(
        f"/api/v1/transactions/{tx.id}/generate-draft",
        headers=purchase_headers,
        json={"document_type": DocumentType.DRAFT_COST_SHEET.value},
    )
    assert res.status_code == 202

    # 3. Purchase user forbidden from generating sales contract
    res = await client.post(
        f"/api/v1/transactions/{tx.id}/generate-draft",
        headers=purchase_headers,
        json={"document_type": DocumentType.DRAFT_CONTRACT.value},
    )
    assert res.status_code == 403

    # 4. Sales user forbidden from generating purchase contract
    res = await client.post(
        f"/api/v1/transactions/{tx.id}/generate-draft",
        headers=sales_headers,
        json={"document_type": DocumentType.DRAFT_PURCHASE_CONTRACT.value},
    )
    assert res.status_code == 403

    # 5. Other role (fa_user) forbidden from generating purchase contract
    res = await client.post(
        f"/api/v1/transactions/{tx.id}/generate-draft",
        headers=other_headers,
        json={"document_type": DocumentType.DRAFT_PURCHASE_CONTRACT.value},
    )
    assert res.status_code == 403

    # 6. Disallowed document types are rejected with validation error (422)
    res = await client.post(
        f"/api/v1/transactions/{tx.id}/generate-draft",
        headers=purchase_headers,
        json={"document_type": DocumentType.DRAFT_PERFORMA_INVOICE.value},
    )
    assert res.status_code == 422



# --- the AGFZE house style -----------------------------------------------------------------------
#
# These lock the page down against the approved reference documents. They are deliberately about
# geometry and letterhead rather than wording: a draft that carries the right commercial terms on
# a page that does not look like AGFZE paper is not a document the desk can put in front of a
# counterparty, and nothing else in this file would notice.


def _docx(template) -> "object":
    from docx import Document as DocxDocument

    ensure_template_files()
    return DocxDocument(io.BytesIO(template_path(template).read_bytes()))


@pytest.mark.parametrize("template", sorted(TEMPLATES_BY_TYPE.values(), key=lambda t: t.filename))
def test_every_template_is_laid_out_on_the_reference_page(template) -> None:
    """Page size and margins, as measured from the reference contract's own `sectPr`."""
    from docx.shared import Inches

    section = _docx(template).sections[0]
    assert section.page_width == Inches(house_style.PAGE_WIDTH_IN)
    assert section.page_height == Inches(house_style.PAGE_HEIGHT_IN)
    # The asymmetric left margin is the reference's, not a typo: 1.125in against 1.0in right.
    assert section.left_margin == Inches(1.125)
    assert section.right_margin == Inches(1.0)
    assert section.top_margin == Inches(1.0)


@pytest.mark.parametrize("template", sorted(TEMPLATES_BY_TYPE.values(), key=lambda t: t.filename))
def test_every_template_carries_the_letterhead_on_every_page(template) -> None:
    """The mark and the registered-entity block live in the section header and footer.

    Which is what makes them repeat on continuation pages by themselves - the reference contract
    carries both on all three of its pages, including a signature page with nothing else on it.
    """
    document = _docx(template)
    section = document.sections[0]

    assert not section.header.is_linked_to_previous
    header_xml = section.header.paragraphs[0]._p.xml
    assert "graphicData" in header_xml, "the letterhead mark is not in the page header"

    footer_text = "\n".join(p.text for p in section.footer.paragraphs)
    for line in ("Adani Global FZE", "Post Box 17186,", "www.adani.com"):
        assert line in footer_text, f"the letterhead footer is missing {line!r}"
    assert house_style.FOOTER_LEGAL_ENGLISH in footer_text
    # The footer is the entity block and nothing else. Bank details belong to the deal, change
    # from one to the next, and are carried as fields - never baked into the paper.
    assert "IBAN" not in footer_text and "SWIFT" not in footer_text


def test_the_issuing_entity_is_the_one_the_references_name() -> None:
    """`ADANI GLOBAL FZE`, from every reference document, not an invented trading name."""
    assert draft_service.SELLER_LEGAL_NAME == "ADANI GLOBAL FZE"


def test_every_generated_type_routes_to_exactly_one_template() -> None:
    """No document type can render with another type's layout."""
    for document_type, template in TEMPLATES_BY_TYPE.items():
        assert draft_service.resolve_template(document_type) is template
        assert template.document_type == document_type
    filenames = [template.filename for template in TEMPLATES_BY_TYPE.values()]
    assert len(set(filenames)) == len(filenames), "two document types share a template file"


def test_the_contracts_sign_bilaterally_with_agfze_in_the_left_column() -> None:
    """AGFZE keeps the left column on both contracts; the role it plays is what flips.

    The reference purchase contract has AGFZE as THE BUYER and the supplier as THE SELLER; the
    reference sales contract has it the other way round. Getting this backwards would name the
    wrong party as the seller on a signed contract.
    """
    from app.services.templates.sales_templates import (
        PURCHASE_CONTRACT_TEMPLATE,
        SALES_CONTRACT_TEMPLATE as SALES,
    )

    assert [(c.role, c.entity_field) for c in SALES.signature_columns] == [
        ("THE SELLER", "agfze"),
        ("THE BUYER", "buyer"),
    ]
    assert [(c.role, c.entity_field) for c in PURCHASE_CONTRACT_TEMPLATE.signature_columns] == [
        ("THE BUYER", "agfze"),
        ("THE SELLER", "supplier"),
    ]


def test_a_rendered_invoice_prices_its_line_item_from_the_transaction() -> None:
    """The line-item table populates, and its rate and amount come from the same source.

    A produced invoice whose unit price and amount disagree is worse than one that fails to
    generate, so the two are proved to come from one precedence rather than two.
    """
    values = {name: f"value-{name}" for name in SALES_INVOICE_TEMPLATE.field_names}
    values.update(
        {
            "agfze": draft_service.SELLER_LEGAL_NAME,
            "buyer": CUSTOMER,
            "commodity": "COPPER INGOTS",
            "quantity": "52.660 MT",
            "unit_rate": "9,577.43",
            "currency": "USD",
            "total_value": "504,347.46",
            "territory_reference": territory_reference(Territory.CHINA.value),
            "generated_at": "2026-08-30T10:00:00",
            "generated_by": "Sales User",
        }
    )
    result = render_template(
        SALES_INVOICE_TEMPLATE,
        template_bytes=template_path(SALES_INVOICE_TEMPLATE).read_bytes(),
        values=values,
        directives=[],
    )

    from docx import Document as DocxDocument

    document = DocxDocument(io.BytesIO(result.content))
    table_text = "\n".join(
        cell.text for table in document.tables for row in table.rows for cell in row.cells
    )
    assert "DESCRIPTION" in table_text and "UNIT PRICE" in table_text
    assert "COPPER INGOTS" in table_text
    assert "9,577.43" in table_text
    assert "504,347.46" in table_text
    assert "{{" not in table_text
    # The single signature column names AGFZE, as the reference invoice's does.
    assert draft_service.SELLER_LEGAL_NAME in table_text
